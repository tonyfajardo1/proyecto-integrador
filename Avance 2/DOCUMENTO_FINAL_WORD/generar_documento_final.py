"""
Script para generar el Documento Final del Proyecto Integrador
Basado en la plantilla de biblioteca USFQ
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Rutas
PLANTILLA = r"F:\proyecto-integrador\Avance 2\guia-biblioteca-pregrado (1).docx"
SALIDA = r"F:\proyecto-integrador\Avance 2\DOCUMENTO_FINAL_WORD\DOCUMENTO_FINAL_AVANCE2.docx"
IMAGENES = r"F:\proyecto-integrador\Avance 2\DOCUMENTO_FINAL_WORD\imagenes"

def agregar_titulo(doc, texto, nivel=1):
    """Agrega un título con el estilo apropiado"""
    if nivel == 1:
        doc.add_heading(texto, level=1)
    elif nivel == 2:
        doc.add_heading(texto, level=2)
    elif nivel == 3:
        doc.add_heading(texto, level=3)

def agregar_parrafo(doc, texto, negrita=False, italica=False):
    """Agrega un párrafo con formato opcional"""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    run.italic = italica
    return p

def agregar_lista(doc, items):
    """Agrega una lista con viñetas"""
    for item in items:
        p = doc.add_paragraph()
        p.add_run("• " + item)

def agregar_imagen(doc, ruta_imagen, ancho_inches=6, pie_figura=""):
    """Agrega una imagen con pie de figura"""
    if os.path.exists(ruta_imagen):
        doc.add_picture(ruta_imagen, width=Inches(ancho_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pie_figura:
            p = doc.add_paragraph(pie_figura)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Imagen no encontrada: {ruta_imagen}]")

def agregar_tabla(doc, encabezados, filas):
    """Agrega una tabla con formato"""
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Table Grid'

    # Encabezados
    hdr_cells = tabla.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        hdr_cells[i].text = encabezado
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Filas
    for fila in filas:
        row_cells = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = str(valor)

    return tabla

def crear_documento():
    """Crea el documento final completo"""

    # Crear documento nuevo
    doc = Document()

    # ===== PORTADA =====
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD SAN FRANCISCO DE QUITO")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Colegio de Ciencias e Ingenierías")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APLICACIÓN DE TÉCNICAS DE DATA MINING PARA ANÁLISIS DE DEVOLUCIONES, DESVIACIONES PLAN VS REAL Y DETECCIÓN DE ANOMALÍAS EN CONDIMENSA")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Proyecto Integrador - Avance 2")
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Anthony Fajardo").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tutor: José Vega")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Marzo 2026")

    doc.add_page_break()

    # ===== RESUMEN =====
    agregar_titulo(doc, "RESUMEN", 1)
    agregar_parrafo(doc,
        "Este proyecto aplica técnicas de Data Mining sobre datos integrados de la empresa CONDIMENSA "
        "para responder tres preguntas analíticas: (1) qué factores permiten pronosticar desviaciones "
        "entre plan y producción real, (2) qué productos se venden juntos para estrategias de cross-selling, "
        "y (3) qué agencias presentan comportamientos atípicos. Se implementó una arquitectura Medallion "
        "(Bronze/Silver/Gold) sobre PostgreSQL, orquestada con Mage AI y visualizada en Streamlit. "
        "Los resultados muestran que ExtraTrees reduce el error de pronóstico en 6.71% vs baseline, "
        "se identificaron 9 reglas de asociación con Lift promedio de 15.57, y se detectó 1 agencia "
        "anómala (QUITO) con estabilidad bootstrap de 70%."
    )

    p = doc.add_paragraph()
    p.add_run("Palabras clave: ").bold = True
    p.add_run("Data Mining, CRISP-DM, Arquitectura Medallion, Pronóstico, Reglas de Asociación, Detección de Anomalías")

    doc.add_page_break()

    # ===== TABLA DE CONTENIDOS =====
    agregar_titulo(doc, "TABLA DE CONTENIDOS", 1)
    contenidos = [
        "1. Introducción",
        "2. Estado del Arte",
        "3. Marco Conceptual y Justificación Arquitectónica",
        "4. Metodología (CRISP-DM)",
        "5. Arquitectura e Implementación",
        "6. Experimentos y Resultados",
        "7. Validación Metodológica",
        "8. Discusión",
        "9. Avance vs Cronograma",
        "10. Conclusiones y Siguientes Pasos",
        "11. Referencias",
        "12. Anexos"
    ]
    for item in contenidos:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ===== 1. INTRODUCCIÓN =====
    agregar_titulo(doc, "1. INTRODUCCIÓN", 1)

    agregar_titulo(doc, "1.1 Contexto", 2)
    agregar_parrafo(doc,
        "CONDIMENSA es una empresa ecuatoriana dedicada a la producción y comercialización de condimentos "
        "y productos alimenticios. La empresa opera con dos sistemas transaccionales principales: "
        "QuickBooks/ODIN para producción y operaciones, y Kronos para gestión comercial. Esta separación "
        "dificulta el análisis integrado de información para la toma de decisiones."
    )

    agregar_titulo(doc, "1.2 Planteamiento del Problema", 2)
    agregar_parrafo(doc,
        "El avance previo priorizó la infraestructura (ETL, capas y dashboard), pero no evidenció con "
        "suficiente rigor la parte de Data Mining. En este avance se corrige ese desbalance, enfocando "
        "la evaluación en preguntas de negocio, métricas adecuadas y validez metodológica."
    )

    agregar_titulo(doc, "1.3 Preguntas Analíticas", 2)
    preguntas = [
        "1. Pronóstico de producción: ¿Qué factores permiten anticipar desviaciones entre cantidad planificada y real?",
        "2. Cross-selling: ¿Qué productos tienden a venderse juntos para promociones combinadas?",
        "3. Anomalías: ¿Qué agencias presentan comportamientos atípicos que requieren investigación?"
    ]
    agregar_lista(doc, preguntas)

    agregar_titulo(doc, "1.4 Objetivo General", 2)
    agregar_parrafo(doc,
        "Aplicar técnicas de Data Mining sobre datos integrados de CONDIMENSA para responder preguntas "
        "analíticas no triviales y generar recomendaciones accionables con evaluación metodológica rigurosa."
    )

    agregar_titulo(doc, "1.5 Alcance y Limitaciones", 2)
    p = doc.add_paragraph("• Período analizado: Julio 2024 - Enero 2026 (18 meses)")
    p = doc.add_paragraph("• Unidades de análisis: producto-período, transacción, agencia")
    p = doc.add_paragraph("• Fuentes: QuickBooks/ODIN, Kronos, PostgreSQL Gold")
    p = doc.add_paragraph("• Limitaciones: N=10 agencias, cobertura de 18 meses")

    doc.add_page_break()

    # ===== 2. ESTADO DEL ARTE =====
    agregar_titulo(doc, "2. ESTADO DEL ARTE", 1)

    agregar_titulo(doc, "2.1 Fundamentos", 2)
    agregar_parrafo(doc,
        "El proyecto sigue CRISP-DM (Cross-Industry Standard Process for Data Mining) para alinear "
        "objetivos de negocio con preparación de datos, modelado y evaluación (Chapman et al., 2000)."
    )

    agregar_titulo(doc, "2.2 Trabajos Relacionados", 2)

    encabezados = ["Trabajo", "Problema", "Técnica", "Métrica", "Aporte"]
    filas = [
        ["Gu et al. (2024)", "Predicción devoluciones", "Random Forest", "AUC-ROC", "Confirma RF para PRB"],
        ["Zhang et al. (2024)", "Gestión devoluciones retail", "LR, RF, GB", "AUC-ROC", "Variables clave"],
        ["Baur et al. (2025)", "Devoluciones textiles", "XGBoost, RF", "Bal. Acc 0.86", "Viabilidad manufactura"],
        ["Santos et al. (2025)", "Anomalías en compras", "K-Means + IF", "Precisión", "Enfoque híbrido"],
        ["Jadhav et al. (2023)", "Market basket", "Apriori", "Lift, Conf", "Cross-selling validado"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_titulo(doc, "2.3 Análisis Crítico", 2)
    agregar_parrafo(doc,
        "La literatura confirma la efectividad de Random Forest y Gradient Boosting para predicción, "
        "con AUC-ROC entre 0.69 y 0.86. Sin embargo, la mayoría de implementaciones no muestran "
        "integración con arquitectura reproducible ni controles anti-leakage explícitos. Este proyecto "
        "contribuye al aplicar estas técnicas en contexto manufacturero ecuatoriano con stack open-source."
    )

    doc.add_page_break()

    # ===== 3. MARCO CONCEPTUAL =====
    agregar_titulo(doc, "3. MARCO CONCEPTUAL Y JUSTIFICACIÓN ARQUITECTÓNICA", 1)

    agregar_titulo(doc, "3.1 OLTP vs OLAP", 2)

    encabezados = ["Aspecto", "OLTP", "OLAP"]
    filas = [
        ["Propósito", "Registrar transacciones", "Analizar históricos"],
        ["Consultas", "Cortas y frecuentes", "Agregadas e intensivas"],
        ["Rendimiento", "Baja latencia", "Alto throughput analítico"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_parrafo(doc,
        "Justificación: Separar OLTP y OLAP evita afectar sistemas operacionales durante análisis "
        "y entrenamiento de modelos."
    )

    agregar_titulo(doc, "3.2 Arquitectura Medallion", 2)
    p = doc.add_paragraph("• Bronze: Datos crudos con mínima transformación")
    p = doc.add_paragraph("• Silver: Datos estandarizados y limpios")
    p = doc.add_paragraph("• Gold: Tablas analíticas para modelado y dashboard")

    doc.add_page_break()

    # ===== 4. METODOLOGÍA =====
    agregar_titulo(doc, "4. METODOLOGÍA (CRISP-DM)", 1)

    agregar_titulo(doc, "4.1 Business Understanding", 2)
    p = doc.add_paragraph("• KPI Pronóstico: WAPE (Weighted Absolute Percentage Error)")
    p = doc.add_paragraph("• KPI Asociación: Lift y Confianza")
    p = doc.add_paragraph("• KPI Anomalías: Estabilidad bootstrap (Jaccard)")

    agregar_titulo(doc, "4.2 Data Understanding", 2)
    p = doc.add_paragraph("• Tablas fuente: bronze.quickbooks_produccion_raw, bronze.kronos_ventas_raw")
    p = doc.add_paragraph("• Cobertura: 18 meses (Jul 2024 - Ene 2026)")
    p = doc.add_paragraph("• Registros: ~10,000 producción, ~2,500 transacciones")

    agregar_titulo(doc, "4.3 Data Preparation", 2)
    p = doc.add_paragraph("• Integración QuickBooks + Kronos en Silver")
    p = doc.add_paragraph("• Features temporales: lag_1, lag_2, lag_3, tendencia_3m")
    p = doc.add_paragraph("• Control de leakage: exclusión de variables post-evento")

    agregar_titulo(doc, "4.4 Modeling", 2)
    p = doc.add_paragraph("• Pronóstico: ExtraTrees, RandomForest, GradientBoosting")
    p = doc.add_paragraph("• Asociación: Apriori, FP-Growth")
    p = doc.add_paragraph("• Anomalías: LOF, Isolation Forest, One-Class SVM")

    agregar_titulo(doc, "4.5 Evaluation", 2)
    agregar_parrafo(doc,
        "Split temporal: Train (Jul24-May25) / Validation (Jun25-Sep25) / Test (Oct25-Ene26). "
        "Métricas: MAE, RMSE, WAPE, Lift, Confianza, Jaccard."
    )

    doc.add_page_break()

    # ===== 5. ARQUITECTURA =====
    agregar_titulo(doc, "5. ARQUITECTURA E IMPLEMENTACIÓN", 1)

    agregar_titulo(doc, "5.1 Arquitectura General", 2)
    agregar_imagen(doc, os.path.join(IMAGENES, "fig01_arquitectura.png"), 6,
        "Figura 1. Arquitectura del Sistema CONDIMENSA: Fuentes → Mage AI (Medallion) → PostgreSQL → Streamlit")

    agregar_titulo(doc, "5.2 Pipelines Implementados", 2)
    agregar_imagen(doc, os.path.join(IMAGENES, "fig02_pipelines.png"), 6,
        "Figura 2. 7 pipelines en Mage AI: 4 Data Mining + 3 ETL")

    encabezados = ["Fuente", "Bronze", "Silver", "Gold", "Data Mining"]
    filas = [
        ["QuickBooks", "etl_bronze", "etl_silver", "etl_gold", "dm_pronostico"],
        ["Kronos", "etl_bronze", "etl_silver", "etl_gold", "dm_asociacion, dm_anomalias"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_imagen(doc, os.path.join(IMAGENES, "fig03_etl_bronze.png"), 6,
        "Figura 3. Ejecución exitosa del pipeline ETL Bronze (26 marzo 2026)")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig12_etl_gold.png"), 6,
        "Figura 4. Ejecución pipeline ETL Gold: cargar, calcular_kpis, extraer, crear_tablas")

    agregar_titulo(doc, "5.3 Reglas de Calidad", 2)
    agregar_imagen(doc, os.path.join(IMAGENES, "fig04_sql_gold.png"), 5,
        "Figura 5. Validación SQL Gold: 2,683 filas totales, 2,656 activas, 27 inactivas")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig19_wrangling.png"), 4,
        "Figura 6. Wrangling Report: 14,670 rows, 904 productos, 21 periodos, 2,434 imputados")

    agregar_titulo(doc, "5.4 Repositorio", 2)
    p = doc.add_paragraph("• GitHub: https://github.com/tonyfajardo1/proyecto-integrador")
    p = doc.add_paragraph("• Estructura: Avance 1/ (base) y Avance 2/ (correcciones)")
    p = doc.add_paragraph("• Semilla: random_state=42")

    doc.add_page_break()

    # ===== 6. EXPERIMENTOS =====
    agregar_titulo(doc, "6. EXPERIMENTOS Y RESULTADOS", 1)

    agregar_imagen(doc, os.path.join(IMAGENES, "fig05_kpis.png"), 6,
        "Figura 7. KPIs: Total Ventas $1,373,944, Rentabilidad $439,613, Tasa Devolución 5.8%, 10 Agencias")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig06_agencias.png"), 6,
        "Figura 8. Distribución ventas por agencia: Quito $512,415 (33%)")

    # 6.1 Pronóstico
    agregar_titulo(doc, "6.1 Experimento 1: Pronóstico de Producción", 2)

    agregar_parrafo(doc, "Objetivo: Predecir cantidad despachada (t+1)", negrita=True)
    agregar_parrafo(doc, "Dataset: 10,316 registros - Train: 5,667 / Validation: 2,432 / Test: 2,217")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig07_benchmark.png"), 6,
        "Figura 9. Benchmark: ExtraTrees (WAPE 0.3774) supera baseline Lag-1 en 6.71%")

    encabezados = ["Modelo", "MAE", "RMSE", "WAPE", "Mejora vs Baseline"]
    filas = [
        ["ExtraTrees", "2,790.86", "9,280.11", "0.3774", "+6.71%"],
        ["RandomForest", "2,860.15", "9,534.26", "0.3867", "+5.77%"],
        ["GradientBoosting", "2,904.00", "9,949.91", "0.3927", "+5.18%"],
        ["Baseline_Lag1", "3,287.20", "11,339.15", "0.4445", "-"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_imagen(doc, os.path.join(IMAGENES, "fig08_pronosticos.png"), 6,
        "Figura 10. Dashboard pronósticos: 724 productos, 19.6M unidades recomendadas")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig13_scatter.png"), 6,
        "Figura 11. Scatter Pronóstico vs Plan histórico con niveles de confianza")

    agregar_parrafo(doc, "Hallazgos:", negrita=True)
    p = doc.add_paragraph("• ExtraTrees reduce error de planificación en 6.71%")
    p = doc.add_paragraph("• Variables importantes: lag_1, lag_2, tendencia_3m")
    p = doc.add_paragraph("• Aplicación: ajustar órdenes de producción con anticipación")

    # 6.2 Cross-Selling
    agregar_titulo(doc, "6.2 Experimento 2: Reglas de Asociación (Cross-Selling)", 2)

    agregar_parrafo(doc, "Objetivo: Identificar productos que se venden juntos", negrita=True)
    agregar_parrafo(doc, "Dataset: ~2,500 transacciones, min_support=0.02, min_confidence=0.35")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig09_apriori.png"), 6,
        "Figura 12. Cross-Selling: 9 reglas, Lift máximo 15.57, Confianza 52%")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig15_apriori_detalle.png"), 6,
        "Figura 13. Interpretación métricas + Lift por regla + Recomendaciones de negocio")

    encabezados = ["Algoritmo", "Reglas", "Lift", "Confianza test", "Jaccard"]
    filas = [
        ["FPGrowth", "9", "15.57", "0.52", "1.00"],
        ["Apriori", "9", "15.57", "0.52", "1.00"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_parrafo(doc,
        "Recomendación principal: Cuando un cliente compra FDA POPULAR 5 NEGRO LLANA, "
        "hay 70% probabilidad de que compre FDA POPULAR 2 NEGRO LLANA (Lift 15.6x).", italica=True)

    agregar_parrafo(doc, "Hallazgos:", negrita=True)
    p = doc.add_paragraph("• 9 reglas estables con confianza >50%")
    p = doc.add_paragraph("• Lift 15.57 indica asociaciones muy fuertes")
    p = doc.add_paragraph("• Aplicación: crear combos, optimizar layout de tienda")

    # 6.3 Anomalías
    agregar_titulo(doc, "6.3 Experimento 3: Detección de Anomalías", 2)

    agregar_parrafo(doc, "Objetivo: Identificar agencias atípicas", negrita=True)
    agregar_parrafo(doc, "Dataset: 10 agencias, contamination=10%")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig10_anomalias.png"), 6,
        "Figura 14. 10 agencias analizadas, 1 anomalía (QUITO)")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig14_anomalias_detalle.png"), 6,
        "Figura 15. Mapa Devolución vs Rentabilidad + Tabla completa de agencias")

    encabezados = ["Algoritmo", "Anomalías", "Bootstrap Jaccard", "Score"]
    filas = [
        ["LOF", "1", "0.70", "0.42"],
        ["IsolationForest", "1", "0.33", "0.20"],
        ["PCA", "1", "0.30", "0.18"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_parrafo(doc, "Detalle QUITO (anomalía):", negrita=True)
    p = doc.add_paragraph("• Tasa Devolución: 5.6%, Rentabilidad: 35.2%, Total Ventas: $512,415")
    p = doc.add_paragraph("• Tipo: ALTA_RENTABILIDAD, BAJO_COSTO, ALTO_TICKET")

    doc.add_page_break()

    # ===== 7. VALIDACIÓN =====
    agregar_titulo(doc, "7. VALIDACIÓN METODOLÓGICA", 1)

    agregar_titulo(doc, "7.1 Control de Leakage", 2)
    p = doc.add_paragraph("• ✓ Target representa evento futuro (t+1)")
    p = doc.add_paragraph("• ✓ Features no contienen información del target")
    p = doc.add_paragraph("• ✓ Split temporal aplicado")
    p = doc.add_paragraph("• ✓ Transformaciones solo en train")

    agregar_titulo(doc, "7.2 Control de Overfitting", 2)
    encabezados = ["Modelo", "Gap Train-Val", "Gap Val-Test", "Evaluación"]
    filas = [
        ["ExtraTrees", "0.2065", "0.0496", "Generaliza bien"]
    ]
    agregar_tabla(doc, encabezados, filas)
    agregar_parrafo(doc, "Gap val-test pequeño (<0.06) confirma generalización adecuada.")

    doc.add_page_break()

    # ===== 8. DISCUSIÓN =====
    agregar_titulo(doc, "8. DISCUSIÓN", 1)

    agregar_titulo(doc, "8.1 Respuestas a Preguntas de Negocio", 2)
    p = doc.add_paragraph("• P1 (Pronóstico): ExtraTrees con WAPE 0.3774 supera baseline en 6.71%")
    p = doc.add_paragraph("• P2 (Cross-selling): 9 reglas estables con Lift 15.57")
    p = doc.add_paragraph("• P3 (Anomalías): QUITO detectada con estabilidad 70%")

    agregar_titulo(doc, "8.2 Implicaciones para CONDIMENSA", 2)
    encabezados = ["Área", "Hallazgo", "Acción", "Impacto"]
    filas = [
        ["Producción", "-6.71% error", "Ajustar órdenes", "Reducir inventario"],
        ["Comercial", "9 reglas", "Crear combos", "Aumentar ticket"],
        ["Control", "1 anomalía", "Auditoría", "Reducir riesgo"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_titulo(doc, "8.3 Limitaciones", 2)
    p = doc.add_paragraph("• N=10 agencias limita robustez de anomalías")
    p = doc.add_paragraph("• 18 meses puede no capturar estacionalidad completa")
    p = doc.add_paragraph("• Validación operativa con stakeholders pendiente")

    doc.add_page_break()

    # ===== 9. AVANCE VS CRONOGRAMA =====
    agregar_titulo(doc, "9. AVANCE VS CRONOGRAMA", 1)

    encabezados = ["Actividad", "Planificado", "Estado"]
    filas = [
        ["Documento planificación", "Semana 1-2", "✓ OK"],
        ["Estado del arte", "Semana 1-3", "✓ OK"],
        ["Data Mart PostgreSQL", "Semana 2-3", "✓ OK"],
        ["Pipelines ETL", "Semana 3-4", "✓ OK"],
        ["Pipelines DM", "Semana 4-5", "✓ OK"],
        ["Dashboard", "Semana 4-5", "✓ OK"],
        ["Experimentación", "Semana 5", "✓ OK"]
    ]
    agregar_tabla(doc, encabezados, filas)

    agregar_parrafo(doc, "Avance: 100% completado", negrita=True)

    doc.add_page_break()

    # ===== 10. CONCLUSIONES =====
    agregar_titulo(doc, "10. CONCLUSIONES Y SIGUIENTES PASOS", 1)

    agregar_titulo(doc, "10.1 Conclusiones", 2)
    p = doc.add_paragraph("• Validez metodológica: Split temporal, control anti-leakage, métricas completas")
    p = doc.add_paragraph("• Arquitectura justificada: OLTP/OLAP con Medallion fundamentada")
    p = doc.add_paragraph("• Resultados accionables: -6.71% error, 9 reglas cross-selling, 1 anomalía")
    p = doc.add_paragraph("• Contribución: Data Mining en manufactura ecuatoriana con stack open-source")

    agregar_titulo(doc, "10.2 Siguientes Pasos", 2)
    p = doc.add_paragraph("• 1. Validación con usuarios CONDIMENSA")
    p = doc.add_paragraph("• 2. Refinamiento de modelos según feedback")
    p = doc.add_paragraph("• 3. Despliegue productivo")
    p = doc.add_paragraph("• 4. Documentación final")

    doc.add_page_break()

    # ===== 11. REFERENCIAS =====
    agregar_titulo(doc, "11. REFERENCIAS", 1)

    referencias = [
        "[1] Gu, Y., Chen, X., & Wang, L. (2024). Understanding and predicting online product return behavior. Int. J. Production Economics, 267, 109066.",
        "[2] Zhang, H., Liu, M., & Chen, S. (2024). Return management in e-commerce firms. J. Cleaner Production, 434, 140123.",
        "[3] Baur, A., Schmidt, K., & Weber, M. (2025). Machine learning for garment returns prediction. SN Computer Science, 6(2), 89.",
        "[4] Santos, R., Ferreira, P., & Oliveira, T. (2025). Anomaly detection in enterprise purchase processes. Information, 16(3), 177.",
        "[5] Jadhav, A., Jadhav, A., & Jadhav, R. D. (2023). Association rule mining in retail. SSRN.",
        "[6] Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit (3rd ed.). Wiley.",
        "[7] Chapman, P. et al. (2000). CRISP-DM 1.0. SPSS Inc.",
        "[8] Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation forest. IEEE ICDM, 413-422.",
        "[9] Agrawal, R., & Srikant, R. (1994). Fast algorithms for mining association rules. VLDB, 487-499."
    ]
    for ref in referencias:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_page_break()

    # ===== 12. ANEXOS =====
    agregar_titulo(doc, "12. ANEXOS", 1)

    agregar_titulo(doc, "Anexo A: Diccionario de Datos", 2)
    agregar_parrafo(doc, "Ver sql/00_crear_todas_tablas.sql en repositorio")

    agregar_titulo(doc, "Anexo B: Evidencias Adicionales", 2)
    agregar_imagen(doc, os.path.join(IMAGENES, "fig16_sql_silver.png"), 5,
        "Figura A1. Query Silver: tipo_producto, producto, periodo, qty_planificada, qty_fabricada")

    agregar_imagen(doc, os.path.join(IMAGENES, "fig21_estructura.png"), 4,
        "Figura A2. Estructura del proyecto de forecasting")

    agregar_titulo(doc, "Anexo C: Repositorio", 2)
    agregar_parrafo(doc, "GitHub: https://github.com/tonyfajardo1/proyecto-integrador")

    # Guardar documento
    doc.save(SALIDA)
    print(f"[OK] Documento guardado en: {SALIDA}")

if __name__ == "__main__":
    crear_documento()
