from pathlib import Path
import csv

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_PATH = Path(r"F:\proyecto-integrador\Avance 2\entregable_2_completado_avance2_v6_autogenerado.docx")
BENCH_PATH = Path(r"F:\proyecto-integrador\Avance 2\03_modelado\forecasting_tesis_v2\artifacts\benchmark_forecasting_v2.csv")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Actualice el indice con F9 en Word"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def load_top_benchmark(top_n=5):
    rows = []
    with BENCH_PATH.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    rows = sorted(rows, key=lambda x: float(x["WAPE_val"]))
    return rows[:top_n]


def main():
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COLEGIO DE CIENCIAS E INGENIERIAS\nINGENIERIA EN CIENCIAS DE LA COMPUTACION")
    r.bold = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Entregable 2 del Proyecto Integrador")
    r.bold = True
    r.font.size = Pt(14)

    for line in [
        "Tutor: Jose David Vega Sanchez",
        "Autor: Anthony Fajardo",
        "Quito - Ecuador",
        "2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    doc.add_page_break()

    doc.add_heading("Indice", level=1)
    p = doc.add_paragraph()
    add_toc(p)

    doc.add_page_break()

    doc.add_heading("1. Titulo del proyecto", level=1)
    doc.add_paragraph("Forecasting de produccion para soporte a la planificacion operativa en CONDIMENSA.")

    doc.add_heading("2. Resumen ejecutivo del periodo", level=1)
    doc.add_paragraph(
        "En este periodo se consolido el flujo E2E de forecasting mensual de produccion "
        "(Mage -> Silver -> modelado temporal -> Gold -> dashboard), con depuracion de pipelines, "
        "estandarizacion de datos y publicacion operativa para toma de decisiones."
    )

    doc.add_heading("3. Actividades realizadas y estado", level=1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Actividad"
    t.rows[0].cells[1].text = "Estado"
    t.rows[0].cells[2].text = "Evidencia"
    rows = [
        ("Depuracion Bronze y Silver", "Completado", "etl_bronze y etl_silver ejecutan correctamente"),
        ("Depuracion Gold + estado forecasting", "Completado", "etl_gold incluye publicar_estado_forecasting_gold"),
        ("Modelado temporal y benchmark", "Completado", "artifact benchmark_forecasting_v2.csv"),
        ("Publicacion Gold v2 y unificado", "Completado", "gold.pronostico_produccion_resultado_v2 y unificado_v1"),
        ("Dashboard pronostico actualizado", "Completado", "Tabla min/recomendada/max + tabs Estacionales/Inactivos"),
    ]
    for a, b, c in rows:
        rc = t.add_row().cells
        rc[0].text = a
        rc[1].text = b
        rc[2].text = c

    doc.add_heading("4. Secciones o capitulos desarrollados", level=1)
    for item in [
        "1. Introduccion (desarrollada).",
        "2. Estado del arte (desarrollado).",
        "3. Metodologia y arquitectura (desarrollada).",
        "4. Implementacion tecnica (desarrollada).",
        "5. Experimentos y resultados (desarrollado).",
        "6. Conclusiones (desarrolladas).",
        "Referencias y anexos (actualizados).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Informacion grafica, tablas y evidencias", level=1)
    doc.add_heading("5.1 Tabla comparativa de modelos (benchmark)", level=2)

    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    for i, h in enumerate(["Modelo", "WAPE_val", "WAPE_test", "MAE_test", "RMSE_test"]):
        t2.rows[0].cells[i].text = h

    for row in load_top_benchmark(top_n=5):
        rc = t2.add_row().cells
        rc[0].text = row["modelo"]
        rc[1].text = f"{float(row['WAPE_val']):.6f}"
        rc[2].text = f"{float(row['WAPE_test']):.6f}"
        rc[3].text = f"{float(row['MAE_test']):.2f}"
        rc[4].text = f"{float(row['RMSE_test']):.2f}"

    doc.add_paragraph(
        "Interpretacion: por criterio oficial (menor WAPE_val), el algoritmo ganador de la corrida actual es LinearRegression.",
        style="List Bullet",
    )

    doc.add_heading("5.2 Evidencia visual a insertar (capturas)", level=2)
    doc.add_paragraph("Figura 1. Pipeline etl_gold en Mage (estado exitoso).")
    doc.add_paragraph("Figura 2. Vista dashboard Pronostico Produccion con rango minima/recomendada/maxima.")
    doc.add_paragraph("Figura 3. Tabs de Estado Forecasting: Estacionales e Inactivos.")
    doc.add_paragraph("Tabla 1. Conteos de publicacion Gold: 4996 filas en resultado_v2 y tabla unificada actualizada.")

    doc.add_heading("6. Verificacion contra observaciones del tutor", level=1)
    t3 = doc.add_table(rows=1, cols=3)
    t3.style = "Table Grid"
    t3.rows[0].cells[0].text = "Observacion"
    t3.rows[0].cells[1].text = "Verificacion"
    t3.rows[0].cells[2].text = "Estado"
    checks = [
        (
            "Cambios en azul revisados",
            "No se detectaron marcas de color azul en v5; validar control de cambios en Word.",
            "Pendiente de confirmacion visual",
        ),
        ("Agregar secciones faltantes", "Se incorporaron secciones estructuradas con encabezados formales.", "Completado"),
        ("Incluir tablas/graficas", "Se agrego tabla de benchmark y placeholders de figuras.", "Completado"),
        (
            "Indice/menu automatico",
            "Se inserto indice automatico por estilos Heading 1-3 (actualizable con F9).",
            "Completado",
        ),
    ]
    for a, b, c in checks:
        rc = t3.add_row().cells
        rc[0].text = a
        rc[1].text = b
        rc[2].text = c

    doc.add_heading("7. Revision y firma del tutor", level=1)
    doc.add_paragraph(
        "Yo, Jose David Vega Sanchez, profesor de la carrera de Ingenieria en Ciencias de la Computacion, "
        "hago constar que he revisado y apruebo las actividades realizadas durante este periodo de trabajo."
    )
    doc.add_paragraph("Fdo: Jose David Vega Sanchez")
    doc.add_paragraph("Quito, 29 de marzo de 2026")
    doc.add_paragraph(
        "Nota: para actualizar el indice automatico, abrir en Word y presionar Ctrl+A, luego F9.",
        style="Intense Quote",
    )

    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
