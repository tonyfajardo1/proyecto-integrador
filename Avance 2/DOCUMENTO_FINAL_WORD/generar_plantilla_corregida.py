"""
Script para generar la Plantilla de Entregables del Proyecto Integrador
Formato basado en plantilla oficial USFQ
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SALIDA = r"F:\proyecto-integrador\Avance 2\DOCUMENTO_FINAL_WORD\ENTREGABLE_2_PLANTILLA_v2.docx"

def agregar_tabla(doc, encabezados, filas):
    """Agrega una tabla con formato"""
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Table Grid'

    hdr_cells = tabla.rows[0].cells
    for i, enc in enumerate(encabezados):
        hdr_cells[i].text = enc
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True

    for fila in filas:
        row_cells = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = str(valor)

    return tabla

def crear_documento():
    """Crea la plantilla de entregables"""

    doc = Document()

    # ============ ENCABEZADO ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD SAN FRANCISCO DE QUITO USFQ")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Colegio de Ciencias e Ingenierias")
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROYECTO INTEGRADOR")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ENTREGABLE 2")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ============ INFORMACION GENERAL ============
    p = doc.add_paragraph()
    run = p.add_run("INFORMACION GENERAL")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    encabezados = ["Campo", "Valor"]
    filas = [
        ["Estudiante", "Anthony Fajardo"],
        ["Codigo", "325019"],
        ["Carrera", "Ingenieria en Ciencias de la Computacion"],
        ["Tutor", "Jose David Vega Sanchez, M.Sc."],
        ["Periodo", "2025-2026"],
        ["Fecha de entrega", "29 de marzo de 2026"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    # ============ TITULO DEL PROYECTO ============
    p = doc.add_paragraph()
    run = p.add_run("TITULO DEL PROYECTO")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_paragraph(
        "Sistema de forecasting de produccion para soporte a la planificacion "
        "operativa en CONDIMENSA"
    )

    doc.add_paragraph()

    # ============ RESUMEN DE ACTIVIDADES ============
    p = doc.add_paragraph()
    run = p.add_run("RESUMEN DE ACTIVIDADES REALIZADAS (Entregable 2)")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Actividades Completadas")
    run.bold = True

    doc.add_paragraph()

    encabezados = ["#", "Actividad", "Descripcion", "Evidencia"]
    filas = [
        ["1", "Correccion metodologica ML", "Split temporal para evitar data leakage", "CHECKLIST_ANTI_LEAKAGE.md"],
        ["2", "Benchmark de algoritmos", "ExtraTrees vs RF vs GB vs baseline. Mejor: ExtraTrees (WAPE 0.3774)", "RESUMEN_BENCHMARK.md"],
        ["3", "Estado del arte", "5 papers: Gu 2024, Zhang 2024, Baur 2025, Santos 2025, Jadhav 2023", "Seccion 2.2 documento"],
        ["4", "Justificacion OLTP/OLAP", "Documentacion de separacion y arquitectura Medallion", "Seccion 3 documento"],
        ["5", "Arquitectura en documento", "Movida de anexos al cuerpo principal (Seccion 5)", "Seccion 5.1 con diagrama"],
        ["6", "Metricas completas", "MAE, RMSE, WAPE, Lift, Confianza, Jaccard", "MATRIZ_METRICAS.md"],
        ["7", "Dashboard Streamlit", "KPIs, pronosticos, reglas, anomalias", "dashboard/"],
        ["8", "Repositorio GitHub", "Commit Avance 2 con estructura organizada", "github.com/tonyfajardo1"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Correcciones Aplicadas segun Retroalimentacion del Entregable 1")
    run.bold = True

    doc.add_paragraph()

    encabezados = ["Observacion del Profesor", "Accion Tomada", "Estado"]
    filas = [
        ["Estado del arte sin trabajos relacionados", "Tabla con 5 papers y analisis critico", "COMPLETADO"],
        ["AUC-ROC = 0.9993 sospechoso", "Split temporal, target t+1, metricas WAPE", "COMPLETADO"],
        ["Arquitectura en anexos", "Movida a Seccion 5 del documento", "COMPLETADO"],
        ["No explica OLTP vs OLAP", "Seccion 3 con justificacion tecnica", "COMPLETADO"],
        ["Falta distribucion de clases", "N=10,316 con split Train/Val/Test", "COMPLETADO"],
        ["Metricas incompletas", "MAE, RMSE, WAPE, Lift, Confianza, Jaccard", "COMPLETADO"],
        ["Crear repositorio GitHub", "Repositorio activo con commits", "COMPLETADO"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    # ============ AVANCE VS CRONOGRAMA ============
    p = doc.add_paragraph()
    run = p.add_run("AVANCE VS CRONOGRAMA")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    encabezados = ["Actividad Planificada", "Fecha Planificada", "Fecha Real", "Estado"]
    filas = [
        ["Documento de planificacion", "Semana 1-2", "Completado", "OK"],
        ["Estado del arte + KPIs", "Semana 1-3", "Completado", "OK"],
        ["Diseno Data Mart PostgreSQL", "Semana 2-3", "Completado", "OK"],
        ["Pipeline 1 (Produccion)", "Semana 3-4", "Completado", "OK"],
        ["Pipeline 2 (Comercial Kronos)", "Semana 3-4", "Completado", "OK"],
        ["Pipeline 3 (Integracion + DM)", "Semana 4-5", "Completado", "OK"],
        ["Dashboard Streamlit", "Semana 4-5", "Completado", "OK"],
        ["Experimentos y documentacion", "Semana 5", "Completado", "OK"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Porcentaje de avance: 100% de actividades planificadas para Entregable 2")
    run.bold = True

    doc.add_paragraph()

    # ============ RESULTADOS PRINCIPALES ============
    p = doc.add_paragraph()
    run = p.add_run("RESULTADOS PRINCIPALES DEL PERIODO")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Resultados Cuantitativos")
    run.bold = True

    doc.add_paragraph()

    encabezados = ["Metrica", "Valor", "Interpretacion"]
    filas = [
        ["WAPE modelo pronostico", "0.3774", "6.71% mejor que baseline"],
        ["Reglas de asociacion", "9 estables", "Lift promedio 15.57"],
        ["Anomalias detectadas", "1 agencia", "QUITO requiere investigacion"],
        ["Productos pronosticados", "724", "Cobertura completa"],
        ["Registros procesados", "14,670", "Dataset completo"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Entregables Generados")
    run.bold = True

    doc.add_paragraph()

    doc.add_paragraph("1. Documento tecnico final (formato biblioteca USFQ)")
    doc.add_paragraph("2. Presentacion de 5 minutos con guion")
    doc.add_paragraph("3. Repositorio GitHub actualizado")
    doc.add_paragraph("4. Dashboard funcional en Streamlit")
    doc.add_paragraph("5. Evidencias de ejecucion (24 capturas)")

    doc.add_paragraph()

    # ============ SIGUIENTES PASOS ============
    p = doc.add_paragraph()
    run = p.add_run("SIGUIENTES PASOS (Entregable 3)")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    doc.add_paragraph("1. Validacion con usuarios de CONDIMENSA")
    doc.add_paragraph("2. Refinamiento de modelos segun feedback operativo")
    doc.add_paragraph("3. Despliegue productivo del dashboard")
    doc.add_paragraph("4. Documentacion final y manual de usuario")

    doc.add_paragraph()
    doc.add_paragraph()

    # ============ FIRMA DEL TUTOR ============
    p = doc.add_paragraph()
    run = p.add_run("FIRMA DEL TUTOR")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    encabezados = ["", ""]
    filas = [
        ["Nombre del Tutor:", "Jose David Vega Sanchez, M.Sc."],
        ["Firma:", "_________________________"],
        ["Fecha:", "____/____/2026"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTA: Este documento debe estar firmado por el tutor para ser valido.")
    run.italic = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Entregables sin firma obtendran calificacion de cero.")
    run.italic = True
    run.bold = True
    run.font.size = Pt(10)

    # Guardar documento
    doc.save(SALIDA)
    print(f"[OK] Plantilla de Entregables guardada en: {SALIDA}")
    print()
    print("IMPORTANTE: Debes obtener la firma del tutor antes de entregar.")
    print("Sin firma = calificacion de cero.")

if __name__ == "__main__":
    crear_documento()
