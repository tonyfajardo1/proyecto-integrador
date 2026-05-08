from pathlib import Path
import csv

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


IN_PATH = Path(r"F:\proyecto-integrador\Avance 2\documento_final_avance2_pulido_v4_APA.docx")
OUT_PATH = Path(r"F:\proyecto-integrador\Avance 2\documento_final_avance2_pulido_v5_APA_autofix.docx")
BENCH_PATH = Path(r"F:\proyecto-integrador\Avance 2\03_modelado\forecasting_tesis_v2\artifacts\benchmark_forecasting_v2.csv")


H1 = {
    "INTRODUCCION",
    "ESTADO DEL ARTE",
    "3. METODOLOGIA Y ARQUITECTURA",
    "4. IMPLEMENTACION TECNICA",
    "5. EXPERIMENTOS Y RESULTADOS",
    "6. CONCLUSIONES",
    "REFERENCIAS",
}

H2 = {
    "3.1 Justificacion OLTP vs OLAP",
    "3.2 Arquitectura Medallion",
    "3.3 Metodologia de trabajo",
    "Arquitectura de la solucion",
    "Pregunta de negocio y criterio de decision",
    "4.1 Pipelines ETL",
    "4.2 Data Warehouse",
    "4.3 Integracion con dashboard",
    "5.1 Experimento 1: Forecasting",
    "5.2 Experimento 2: Cross-Selling",
    "5.3 Experimento 3: Anomalias",
    "Resultados integrados del avance 2",
    "6.1 Principales hallazgos",
    "6.2 Trabajo futuro",
    "6.3 Limitaciones y consideraciones",
}


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    txt = OxmlElement("w:t")
    txt.text = "Actualice la tabla de contenido con F9"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def load_top5():
    rows = []
    if not BENCH_PATH.exists():
        return rows
    with BENCH_PATH.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)
    rows.sort(key=lambda r: float(r["WAPE_val"]))
    return rows[:5]


def main():
    doc = Document(IN_PATH)

    toc_idx = None
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t in H1:
            p.style = "Heading 1"
        elif t in H2:
            p.style = "Heading 2"

        if t == "TABLA DE CONTENIDO":
            toc_idx = i

    if toc_idx is not None:
        # Limpiar lineas de indice manual con puntos.
        for j in range(toc_idx + 1, min(toc_idx + 20, len(doc.paragraphs))):
            t = (doc.paragraphs[j].text or "").strip()
            if "..." in t or "……" in t:
                clear_paragraph(doc.paragraphs[j])

        # Insertar TOC automatico justo despues del titulo.
        p_toc = doc.paragraphs[toc_idx + 1]
        if not (p_toc.text or "").strip():
            add_toc(p_toc)
        else:
            p_toc.insert_paragraph_before("")
            add_toc(doc.paragraphs[toc_idx + 1])

    # Agregar tabla real de benchmark al final de 5.1 para cumplir evidencia tabular.
    top = load_top5()
    if top:
        doc.add_page_break()
        doc.add_heading("ANEXO A. Tabla de benchmark (auto-generada)", level=1)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        headers = ["Modelo", "WAPE_val", "WAPE_test", "MAE_test", "RMSE_test"]
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h

        for row in top:
            c = tbl.add_row().cells
            c[0].text = row["modelo"]
            c[1].text = f"{float(row['WAPE_val']):.6f}"
            c[2].text = f"{float(row['WAPE_test']):.6f}"
            c[3].text = f"{float(row['MAE_test']):.2f}"
            c[4].text = f"{float(row['RMSE_test']):.2f}"

    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
