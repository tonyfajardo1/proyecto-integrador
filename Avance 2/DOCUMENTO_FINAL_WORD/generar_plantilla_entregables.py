"""
Script para generar la Plantilla de Entregables del Proyecto Integrador
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

# Rutas
PLANTILLA = r"F:\proyecto-integrador\Avance 2\plantilla_entregables (1).docx"
SALIDA = r"F:\proyecto-integrador\Avance 2\DOCUMENTO_FINAL_WORD\ENTREGABLE_2_PLANTILLA_ACTIVIDADES.docx"

def agregar_tabla(doc, encabezados, filas, ancho_columnas=None):
    """Agrega una tabla con formato"""
    tabla = doc.add_table(rows=1, cols=len(encabezados))
    tabla.style = 'Table Grid'

    # Encabezados
    hdr_cells = tabla.rows[0].cells
    for i, encabezado in enumerate(encabezados):
        hdr_cells[i].text = encabezado
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Filas
    for fila in filas:
        row_cells = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = str(valor)

    return tabla

def crear_documento():
    """Crea la plantilla de entregables"""

    # Crear documento nuevo
    doc = Document()

    # ===== ENCABEZADO =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD SAN FRANCISCO DE QUITO")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROYECTO INTEGRADOR - ENTREGABLE 2")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== INFORMACIÓN GENERAL =====
    p = doc.add_paragraph()
    p.add_run("1. INFORMACIÓN GENERAL").bold = True

    doc.add_paragraph()

    encabezados = ["Campo", "Valor"]
    filas = [
        ["Estudiante", "Anthony Fajardo"],
        ["Código", "[Completar código]"],
        ["Carrera", "Ingeniería en Ciencias de la Computación"],
        ["Tutor", "José Vega"],
        ["Período", "2025-2026"],
        ["Fecha de entrega", "29 de marzo de 2026"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    # ===== TÍTULO DEL PROYECTO =====
    p = doc.add_paragraph()
    p.add_run("2. TÍTULO DEL PROYECTO").bold = True

    doc.add_paragraph(
        "Aplicación de técnicas de Data Mining para análisis de devoluciones, "
        "desviaciones plan vs real y detección de anomalías en CONDIMENSA"
    )

    doc.add_paragraph()

    # ===== RESUMEN DE ACTIVIDADES =====
    p = doc.add_paragraph()
    p.add_run("3. RESUMEN DE ACTIVIDADES REALIZADAS (Período Avance 2)").bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("3.1 Actividades Completadas").bold = True

    encabezados = ["#", "Actividad", "Descripción", "Evidencia"]
    filas = [
        ["1", "Corrección metodológica ML", "Split temporal para evitar data leakage", "CHECKLIST_ANTI_LEAKAGE.md"],
        ["2", "Benchmark de algoritmos", "ExtraTrees vs RF vs GB vs baseline. Mejor: ExtraTrees (WAPE 0.3774)", "RESUMEN_BENCHMARK_FORECASTING.md"],
        ["3", "Estado del arte", "5 papers: Gu 2024, Zhang 2024, Baur 2025, Santos 2025, Jadhav 2023", "Sección 2.2 documento"],
        ["4", "Justificación OLTP/OLAP", "Documentación de separación y arquitectura Medallion", "Sección 3 documento"],
        ["5", "Arquitectura en documento", "Movida de anexos al cuerpo principal (Sección 5)", "Sección 5.1 con diagrama"],
        ["6", "Métricas completas", "MAE, RMSE, WAPE, Lift, Confianza, Jaccard", "MATRIZ_METRICAS.md"],
        ["7", "Dashboard Streamlit", "KPIs, pronósticos, reglas, anomalías", "dashboard/"],
        ["8", "Repositorio GitHub", "Commit Avance 2 con estructura organizada", "github.com/tonyfajardo1"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("3.2 Correcciones Aplicadas según Retroalimentación del Entregable 1").bold = True

    encabezados = ["Observación del Profesor", "Acción Tomada", "Estado"]
    filas = [
        ["Estado del arte sin trabajos relacionados", "Tabla con 5 papers y análisis crítico", "✓ COMPLETADO"],
        ["AUC-ROC = 0.9993 sospechoso", "Split temporal, target t+1, métricas WAPE", "✓ COMPLETADO"],
        ["Arquitectura en anexos", "Movida a Sección 5 del documento", "✓ COMPLETADO"],
        ["No explica OLTP vs OLAP", "Sección 3 con justificación técnica", "✓ COMPLETADO"],
        ["Falta distribución de clases", "N=10,316 con split Train/Val/Test", "✓ COMPLETADO"],
        ["Métricas incompletas", "MAE, RMSE, WAPE, Lift, Confianza, Jaccard", "✓ COMPLETADO"],
        ["Crear repositorio GitHub", "Repositorio activo con commits", "✓ COMPLETADO"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    # ===== AVANCE VS CRONOGRAMA =====
    p = doc.add_paragraph()
    p.add_run("4. AVANCE VS CRONOGRAMA").bold = True

    doc.add_paragraph()

    encabezados = ["Actividad Planificada", "Fecha Planificada", "Fecha Real", "Estado"]
    filas = [
        ["Documento de planificación", "Semana 1-2", "Completado", "✓ OK"],
        ["Estado del arte + KPIs", "Semana 1-3", "Completado", "✓ OK"],
        ["Diseño Data Mart PostgreSQL", "Semana 2-3", "Completado", "✓ OK"],
        ["Pipeline 1 (Producción)", "Semana 3-4", "Completado", "✓ OK"],
        ["Pipeline 2 (Comercial Kronos)", "Semana 3-4", "Completado", "✓ OK"],
        ["Pipeline 3 (Integración + DM)", "Semana 4-5", "Completado", "✓ OK"],
        ["Dashboard Streamlit", "Semana 4-5", "Completado", "✓ OK"],
        ["Experimentos y documentación", "Semana 5", "Completado", "✓ OK"]
    ]
    agregar_tabla(doc, encabezados, filas)

    p = doc.add_paragraph()
    p.add_run("Porcentaje de avance: 100% de actividades planificadas para Avance 2").bold = True

    doc.add_paragraph()

    # ===== RESULTADOS PRINCIPALES =====
    p = doc.add_paragraph()
    p.add_run("5. RESULTADOS PRINCIPALES DEL PERÍODO").bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("5.1 Resultados Cuantitativos").bold = True

    encabezados = ["Métrica", "Valor", "Interpretación"]
    filas = [
        ["WAPE modelo pronóstico", "0.3774", "6.71% mejor que baseline"],
        ["Reglas de asociación", "9 estables", "Lift promedio 15.57"],
        ["Anomalías detectadas", "1 agencia", "QUITO requiere investigación"],
        ["Productos pronosticados", "724", "Cobertura completa"],
        ["Registros procesados", "14,670", "Dataset completo"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("5.2 Entregables Generados").bold = True

    doc.add_paragraph("1. Documento técnico final (formato biblioteca USFQ)")
    doc.add_paragraph("2. Presentación de 5 minutos con guion")
    doc.add_paragraph("3. Repositorio GitHub actualizado")
    doc.add_paragraph("4. Dashboard funcional en Streamlit")
    doc.add_paragraph("5. Evidencias de ejecución (24 capturas)")

    doc.add_paragraph()

    # ===== SIGUIENTES PASOS =====
    p = doc.add_paragraph()
    p.add_run("6. SIGUIENTES PASOS (Avance 3)").bold = True

    doc.add_paragraph("1. Validación con usuarios de CONDIMENSA")
    doc.add_paragraph("2. Refinamiento de modelos según feedback operativo")
    doc.add_paragraph("3. Despliegue productivo del dashboard")
    doc.add_paragraph("4. Documentación final y manual de usuario")

    doc.add_paragraph()
    doc.add_paragraph()

    # ===== FIRMA DEL TUTOR =====
    p = doc.add_paragraph()
    p.add_run("7. FIRMA DEL TUTOR").bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    encabezados = ["", ""]
    filas = [
        ["Nombre del Tutor:", "José Vega"],
        ["Firma:", "_________________________"],
        ["Fecha:", "____/____/2026"]
    ]
    agregar_tabla(doc, encabezados, filas)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NOTA: Este documento debe estar firmado por el tutor para ser válido.")
    run.italic = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Entregables sin firma obtendrán calificación de cero.")
    run.italic = True
    run.bold = True
    run.font.size = Pt(10)

    # Guardar documento
    doc.save(SALIDA)
    print(f"[OK] Plantilla guardada en: {SALIDA}")

if __name__ == "__main__":
    crear_documento()
