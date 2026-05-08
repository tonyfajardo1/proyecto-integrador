from docx import Document


SRC = r"F:\proyecto-integrador\Avance 2\documento_final_avance2_pulido_v3_APA.docx"
OUT = r"F:\proyecto-integrador\Avance 2\documento_final_avance2_pulido_v4_APA.docx"


def main():
    d = Document(SRC)

    updates = {
        102: "1. Introduccion.............................................1",
        103: "2. Estado del arte..........................................3",
        104: "3. Metodologia y Arquitectura...............................6",
        105: "   3.1 Justificacion OLTP vs OLAP...........................6",
        106: "   3.2 Arquitectura Medallion...............................7",
        107: "4. Implementacion Tecnica...................................9",
        108: "   4.1 Pipelines ETL........................................9",
        109: "   4.2 Data Warehouse.......................................10",
        110: "5. Experimentos y Resultados................................12",
        111: "   5.1 Experimento 1: Forecasting...........................12",
        112: (
            "   5.2 Experimento 2: Cross-Selling.........................14\n"
            "   5.3 Experimento 3: Anomalias.............................15\n"
            "6. Conclusiones.............................................17\n"
            "Referencias.................................................19\n"
            "Anexos......................................................21"
        ),
        145: "3. METODOLOGIA Y ARQUITECTURA",
        147: "3.1 Justificacion OLTP vs OLAP",
        148: (
            "La propuesta adopta una separacion explicita entre sistemas transaccionales "
            "(OLTP) y analiticos (OLAP). Esta decision responde a criterios tecnicos y de "
            "negocio: i) evitar que consultas analiticas degraden el rendimiento operacional, "
            "ii) preservar trazabilidad historica para comparaciones mensuales, y iii) asegurar "
            "reproducibilidad de modelos y reportes. En OLTP se priorizan operaciones de "
            "escritura/actualizacion y consistencia operacional; en OLAP se priorizan agregaciones, "
            "consultas historicas y explotacion analitica."
        ),
        149: (
            "En el contexto de CONDIMENSA, esta separacion permite que la operacion diaria "
            "continue sin impacto por cargas de analitica y que la evaluacion del forecasting "
            "use datos consolidados y estables. Ademas, facilita auditar de donde proviene "
            "cada KPI y cada recomendacion publicada en dashboard."
        ),
        150: "3.2 Arquitectura Medallion",
        151: (
            "Se implementa arquitectura Medallion con tres capas: Bronze (datos crudos de origen), "
            "Silver (estandarizacion y calidad), y Gold (salidas analiticas y de consumo). Bronze "
            "conserva granularidad y evidencia de ingesta; Silver corrige tipos, limpia inconsistencias "
            "y consolida estructuras para modelado; Gold publica resultados listos para decision."
        ),
        152: (
            "En Silver se concentra la logica de preparacion del dataset mensual (producto, periodo, "
            "qty_planificada, qty_fabricada, n_ordenes), incluyendo control de nulos, duplicados y "
            "normalizacion de tipos de producto (PT/PP/OTRO). En Gold se publica el resultado operativo "
            "en gold.pronostico_produccion_resultado_v2 con campos de trazabilidad (pipeline_id, "
            "fecha_ejecucion, modelo_ganador)."
        ),
        153: "3.3 Metodologia de trabajo",
        154: (
            "La metodologia combina ingenieria de datos y modelado temporal. Primero, se realiza EDA "
            "para validar cobertura y consistencia. Segundo, se ejecuta wrangling temporal con completitud "
            "mensual por SKU, imputacion jerarquica para faltantes y tratamiento inteligente de outliers "
            "sospechosos. Tercero, se entrena y evalua un benchmark de modelos con split temporal estricto "
            "y controles anti-leakage."
        ),
        155: (
            "Como regla de validez, todas las transformaciones que pueden inducir fuga se calibran con "
            "informacion de entrenamiento (o train+val cuando corresponde) y se valida su comportamiento "
            "en test separado temporalmente. Finalmente, los resultados se publican en Gold y se consumen "
            "en dashboard para soporte de decisiones."
        ),
        156: "",
        167: "4. IMPLEMENTACION TECNICA",
        169: "4.1 Pipelines ETL",
        170: (
            "La orquestacion se implementa en Mage con pipelines de Bronze, Silver y Gold. En Bronze "
            "se extraen datos de fuentes operacionales y se registran evidencias de carga. En Silver se "
            "realizan validaciones de calidad, estandarizacion de campos y consolidacion mensual de datos "
            "para modelado."
        ),
        171: (
            "En la etapa de modelado (forecasting_tesis_v2) se ejecutan scripts reproducibles para EDA, "
            "wrangling, entrenamiento y generacion de artefactos. Posteriormente, un script de publicacion "
            "inserta resultados en la tabla Gold v2 que consume el dashboard."
        ),
        172: "4.2 Data Warehouse",
        173: (
            "El Data Warehouse en PostgreSQL separa funciones por capa. Silver contiene tablas de trabajo "
            "curadas para analitica, mientras Gold contiene tablas de salida orientadas a consumo. Esta "
            "estructura facilita gobierno de datos, control de versiones de pipeline y auditoria de resultados."
        ),
        174: (
            "La tabla gold.pronostico_produccion_resultado_v2 integra informacion historica y pronostico: "
            "producto, periodo, pronostico_qty, qty_planificada, recomendaciones, nivel_confianza, vigencia "
            "operativa y metadatos de ejecucion. Esta salida permite conectar modelado con decision operativa "
            "en una misma vista."
        ),
        175: "4.3 Integracion con dashboard",
        176: (
            "El dashboard Streamlit consume primero la tabla Gold v2 y mantiene fallback a version legacy "
            "para continuidad operativa. Se incluyeron KPIs de planificacion, alertas de revision manual y "
            "tabla de productos excluidos por vigencia. Esta integracion cierra el ciclo desde dato crudo "
            "hasta decision visible para usuarios."
        ),
        178: "5. EXPERIMENTOS Y RESULTADOS",
        180: "5.1 Experimento 1: Forecasting",
        181: (
            "Se evaluaron baselines temporales (Lag1, Lag12 estacional, baseline hibrido) y modelos de ML "
            "(RandomForest, ExtraTrees, LinearRegression, Prophet y variantes segmentadas). La evaluacion se "
            "hizo con split temporal train/val/test y metricas WAPE, MAE y RMSE. Resultado principal: "
            "RandomForest obtuvo mejor desempeno por validacion (WAPE_val=0.3637) y mantuvo estabilidad "
            "en test (WAPE_test=0.4102), superando baselines simples."
        ),
        182: "5.2 Experimento 2: Cross-Selling",
        183: (
            "Como analisis complementario (no eje principal del avance 2), se mantuvo la linea de reglas "
            "de asociacion iniciada en avances previos para identificar combinaciones frecuentes de compra. "
            "Este experimento se utiliza como soporte comercial y no como criterio de planificacion de "
            "produccion. Se considera linea secundaria para trabajo futuro en integracion comercial-operativa."
        ),
        184: "5.3 Experimento 3: Anomalias",
        185: (
            "Tambien se mantuvo la linea de deteccion de anomalias iniciada previamente para identificar "
            "comportamientos atipicos en indicadores operativos/comerciales. En el alcance actual, este "
            "experimento tiene rol de alerta complementaria, mientras que el objetivo central y evaluado "
            "en profundidad corresponde al forecasting mensual de produccion."
        ),
        186: "Resultados integrados del avance 2",
        187: (
            "Indicadores clave: wrangling con 14,670 filas y 904 productos; 2,434 meses imputados; 37 "
            "productos estacionales detectados; 21 caps de outliers sospechosos aplicados. En comparacion "
            "formal contra plan humano, el modelo no supera al plan en este corte (WAPE_modelo=0.3940 "
            "vs WAPE_planificada=0.0980), por lo que se adopta estrategia hibrida operacional."
        ),
        189: "6. CONCLUSIONES",
        191: "6.1 Principales hallazgos",
        192: (
            "El proyecto consolido un flujo E2E defendible para forecasting: arquitectura de datos funcional, "
            "calidad trazable, evaluacion temporal rigurosa y despliegue operativo en dashboard. Se demostro "
            "mejora del modelo frente a baselines temporales, y a la vez se documento con transparencia que "
            "la planificacion humana actual mantiene menor error en este corte."
        ),
        194: "6.2 Trabajo futuro",
        195: (
            "Se propone fortalecer cuatro frentes: i) comparacion contra plan ex-ante para eliminar sesgo por "
            "ajustes posteriores, ii) incorporacion de variables exogenas (promociones, quiebres, calendario "
            "comercial), iii) monitoreo de drift por segmento, y iv) evolucion de reglas hibridas hacia politicas "
            "automatizadas por categoria/tipo."
        ),
        197: "6.3 Limitaciones y consideraciones",
        198: (
            "La principal limitacion del corte actual es la alta correlacion entre plan humano y valor real observado, "
            "lo que reduce el margen de mejora del modelo en evaluacion comparativa. Pese a ello, el sistema aporta "
            "valor en gobernanza, trazabilidad, alerta temprana y estandarizacion de decisiones analiticas."
        ),
        199: "REFERENCIAS",
        214: "ANEXOS",
        215: "Anexo A: Evidencias de pipelines y consultas SQL",
        218: "Anexo B: Evidencias de dashboard operativo",
        226: "Anexo C: Artefactos y reportes de modelado",
    }

    for idx, txt in updates.items():
        if idx < len(d.paragraphs):
            d.paragraphs[idx].text = txt

    d.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
