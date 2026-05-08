from docx import Document


SRC = r"F:\proyecto-integrador\Avance 1\Entregable_1_Documento_Final_USFQ.docx"
OUT = r"F:\proyecto-integrador\Avance 2\documento_final_avance2_pulido_v1.docx"


def main():
    d = Document(SRC)

    repl = {
        9: "Sistema de forecasting de produccion para soporte a la planificacion operativa en CONDIMENSA",
        30: "Quito, 29 de marzo de 2026",
        41: "Sistema de forecasting de produccion para soporte a la planificacion operativa en CONDIMENSA",
        47: "Quito, 29 de marzo de 2026",
        62: "Lugar y fecha:                   Quito, 29 de marzo de 2026",
        86: (
            "Este trabajo presenta la implementacion de un sistema de forecasting mensual de produccion para "
            "CONDIMENSA, orientado a apoyar la planificacion operativa con evidencia cuantitativa y trazabilidad "
            "de datos. Se desarrollo un flujo end-to-end que integra orquestacion ETL con Mage, consolidacion "
            "analitica en arquitectura Medallion (Bronze-Silver-Gold), preparacion temporal de datos y evaluacion "
            "de modelos de pronostico.\n\n"
            "El dataset base de modelado proviene de silver.produccion_modelado_mensual. En la fase de calidad se "
            "validaron nulos, duplicados producto-periodo, consistencia de tipos y cobertura temporal por SKU. El "
            "wrangling incluyo completitud mensual, imputacion jerarquica para faltantes (cero estructural en "
            "estacionales, imputacion temporal y mediana de respaldo), clasificacion robusta PT/PP/OTRO y "
            "tratamiento inteligente de outliers sospechosos con trazabilidad.\n\n"
            "Se evaluaron baselines y modelos con validacion temporal estricta y controles anti-leakage. El mejor "
            "modelo por validacion fue RandomForest (WAPE_val=0.3637; WAPE_test=0.4102), superando a baselines "
            "temporales. Sin embargo, en comparacion formal contra la planificacion humana (qty_planificada), el "
            "plan actual obtuvo menor error en este corte (WAPE_planificada=0.0980 vs WAPE_modelo=0.3940), por lo "
            "que se adopto una estrategia hibrida: plan humano como base y modelo como sistema de alertas para "
            "revision de SKU criticos.\n\n"
            "Los resultados se publicaron en gold.pronostico_produccion_resultado_v2 y se integraron al dashboard "
            "Streamlit con trazabilidad de ejecucion, seccion de productos inactivos y alertas de revision manual."
        ),
        90: "Palabras clave: forecasting, series de tiempo, arquitectura medallion, ETL, calidad de datos, leakage, planificacion de produccion, dashboard, machine learning.",
        94: (
            "This project presents the implementation of a monthly production forecasting system for CONDIMENSA "
            "to support operational planning through a reproducible and auditable analytics pipeline. The solution "
            "integrates ETL orchestration with Mage, Medallion architecture (Bronze-Silver-Gold), time-series data "
            "preparation, and model benchmarking under strict temporal validation.\n\n"
            "The modeling dataset is sourced from silver.produccion_modelado_mensual. Data-quality controls include "
            "null checks, product-period duplicate validation, temporal coverage checks, robust PT/PP/OTHER product-type "
            "assignment, hierarchical imputation for missing months, and intelligent treatment of suspicious outliers "
            "with full traceability.\n\n"
            "Forecasting baselines and machine-learning models were compared using anti-leakage safeguards. RandomForest "
            "achieved the best validation performance (WAPE_val=0.3637; WAPE_test=0.4102), outperforming temporal baselines. "
            "However, a formal comparison against human planning (qty_planificada) showed lower error for the current planning "
            "process in this data cut (WAPE_planificada=0.0980 vs WAPE_modelo=0.3940). Therefore, a hybrid operational strategy "
            "was adopted: keep human planning as baseline and use the model as an alerting layer for high-risk SKUs.\n\n"
            "Outputs are published to gold.pronostico_produccion_resultado_v2 and consumed in a Streamlit dashboard with execution "
            "traceability, inactive-product view, and manual-review alerts."
        ),
        98: "Key words: forecasting, time series, Medallion architecture, ETL, data quality, leakage, production planning, dashboard, machine learning.",
        102: "Introduccion................................................1",
        103: "Estado del arte.............................................3",
        104: "Metodologia y arquitectura..................................6",
        105: "Implementacion tecnica y modelado...........................9",
        106: "Resultados y discusion.....................................13",
        107: "Conclusiones y trabajo futuro..............................16",
        108: "Referencias bibliograficas.................................18",
        109: "Anexo A: Evidencias de pipelines y SQL.....................20",
        110: "Anexo B: Evidencias de dashboard...........................22",
        111: "Anexo C: Artefactos y reportes de modelado.................24",
        115: "CONDIMENSA requiere mejorar la planificacion de produccion mensual en un contexto donde los datos operativos historicamente han estado distribuidos en sistemas transaccionales distintos. Esta situacion dificulta comparar rapidamente plan vs real, detectar desviaciones tempranas y sostener decisiones operativas con evidencia reproducible.",
        116: "El proyecto reorienta el alcance hacia forecasting de produccion como problema principal de negocio. Para ello se separa claramente la capa transaccional (OLTP) de la capa analitica (OLAP), implementando arquitectura Medallion para estandarizar, curar y publicar datos listos para modelado y consumo en dashboard.",
        117: "El aporte principal no es solo un modelo predictivo, sino un proceso completo y defendible: calidad de datos, trazabilidad de transformaciones, evaluacion temporal sin leakage, comparacion formal con la planificacion humana y despliegue operativo en Gold y Streamlit.",
        118: "La pregunta de negocio central es: \"Que cantidad debe planificarse por producto para el siguiente mes, reduciendo riesgo operativo y manteniendo trazabilidad metodologica?\". Esta pregunta se responde con un pipeline de series temporales y reglas de decision hibridas para uso real de planificacion.",
        119: "El documento se organiza en seis bloques: estado del arte, metodologia y arquitectura, implementacion tecnica, resultados y discusion, conclusiones y lineas futuras, y anexos de evidencia para validacion academica y tecnica.",
        120: "",
        121: "",
        125: "Fundamentos y trabajos relacionados",
        126: "La literatura de forecasting en entornos operativos enfatiza tres pilares: calidad de datos, evaluacion temporal correcta y uso de metricas alineadas al impacto de negocio. En series de produccion, errores de preparacion (faltantes, estacionalidad mal tratada, leakage) pueden inflar artificialmente el desempeno y generar decisiones no confiables en despliegue real.",
        127: "Diversos trabajos reportan que modelos de arboles y ensambles suelen comportarse de forma robusta en contextos tabulares con senales temporales, siempre que exista validacion por ventanas de tiempo y control de fuga de informacion. Por ello, este proyecto prioriza WAPE/MAE/RMSE en splits temporales y evita evaluaciones aleatorias.",
        129: "OLTP, OLAP y arquitectura Medallion",
        130: "OLTP (Online Transaction Processing) se orienta a registrar transacciones del negocio con alta concurrencia y baja latencia. OLAP (Online Analytical Processing) se orienta a analisis historico y consultas agregadas. Separar ambas capas evita degradar sistemas operacionales y mejora la gobernanza de datos para analitica.",
        131: "La arquitectura Medallion operacionaliza esta separacion: Bronze conserva la ingesta cruda, Silver estandariza y valida, y Gold publica tablas analiticas listas para consumo. Esta estructura permite trazabilidad, reproducibilidad y control de calidad por etapa.",
        133: "Evaluacion temporal y leakage",
        134: "En forecasting, la validez metodologica depende de evitar leakage. Se aplicaron controles explicitos: target desplazado por horizonte (t+1), split temporal estricto train/val/test, encoding de identificadores entrenado solo con train, reglas estacionales definidas con train+val, y caps de prediccion por historico para estabilidad operativa.",
        136: "Criterios de evaluacion",
        137: "Se adopto WAPE como metrica principal por su interpretabilidad operacional sobre volumenes heterogeneos, complementada con MAE y RMSE. Adicionalmente se realizo comparacion formal contra qty_planificada para medir valor real frente a la practica vigente.",
        139: "Trabajo relacionado en contexto del proyecto",
        140: "El proyecto responde directamente a observaciones del primer avance: la arquitectura se integra en el desarrollo metodologico y no en anexos; el estado del arte se centra en trabajos relacionados; y la evaluacion evita resultados artificialmente altos por fuga de informacion o particiones inapropiadas.",
        142: "",
        143: "",
        147: "Objetivo general",
        148: "Implementar un sistema de forecasting mensual de produccion para CONDIMENSA que permita recomendar cantidades por producto mediante un flujo reproducible de datos y modelado temporal, con despliegue operativo en dashboard.",
        150: "Objetivos especificos",
        151: "1. Consolidar datos de produccion en arquitectura Medallion con controles de calidad y trazabilidad de ejecucion.",
        152: "2. Preparar un dataset temporal robusto (EDA, completitud mensual, estacionalidad, imputacion y control de outliers).",
        153: "3. Entrenar y comparar baselines y modelos de forecasting con validacion temporal estricta y controles anti-leakage.",
        154: "4. Publicar predicciones y recomendaciones en Gold v2 para consumo operativo en dashboard.",
        155: "5. Comparar formalmente el desempeno del modelo frente a qty_planificada y definir una estrategia de uso real.",
        156: "6. Documentar evidencia tecnica (artefactos, capturas, reportes) para sustentacion academica.",
        157: "",
        159: "Arquitectura de la solucion",
        160: "La solucion integra: (i) fuentes operativas de produccion/comercial, (ii) orquestacion ETL con Mage, (iii) almacenamiento analitico en PostgreSQL con Bronze/Silver/Gold, (iv) capa de modelado en Python (forecasting_tesis_v2), y (v) dashboard Streamlit para visualizacion y soporte de decisiones.",
        162: "Pregunta de negocio y criterio de decision",
        163: "Pregunta central: como planificar produccion mensual por producto con menor error y mayor trazabilidad operativa?",
        164: "Criterio de decision aplicado: usar modelo cuando aporta valor validado; mantener plan humano como base cuando su desempeno es superior; y activar alertas para revision manual en brechas relevantes.",
        165: "",
        169: "Diseno de datos y wrangling temporal",
        170: "El dataset de modelado se alimenta de silver.produccion_modelado_mensual. Se aplicaron controles de nulos y duplicados, normalizacion semantica de productos, clasificacion PT/PP/OTRO, completitud mensual por SKU, imputacion jerarquica y tratamiento inteligente de outliers sospechosos con reporte de caps aplicados.",
        172: "Modelado y benchmark",
        173: "Se evaluaron Baseline_Lag1, Baseline_Lag12_Seasonal, Baseline_Hibrido_L1_L12, RandomForest, ExtraTrees, LinearRegression, Prophet y variantes segmentadas (por tipo, categoria y jerarquico). Se selecciono el modelo por WAPE de validacion temporal y se verifico estabilidad en test.",
        175: "Publicacion y dashboard",
        176: "La salida se publica en gold.pronostico_produccion_resultado_v2 con campos de trazabilidad (pipeline_id, fecha_ejecucion, modelo_ganador), vigencia operativa e indicadores para planificacion. El dashboard consume v2 con fallback a legado, muestra alertas de revision manual y tabla de productos inactivos.",
        181: "Stack tecnologico",
        184: "Se utiliza stack open-source: Mage AI, PostgreSQL, Python, scikit-learn, Prophet (opcional), Streamlit, Plotly y Docker. La ejecucion del pipeline genera artefactos CSV para auditoria tecnica (benchmark, leakage_report, wrangling_report, plan_vs_model, segment_error_report).",
        186: "Resultados principales del avance 2",
        187: "- Wrangling: 14,670 filas, 904 productos, 21 periodos; 2,434 meses imputados; 37 productos estacionales detectados.\n- Control de outliers: 108 extremos en qty_fabricada, 117 en qty_planificada, 21 caps aplicados sobre casos sospechosos.\n- Mejor modelo por validacion: RandomForest (WAPE_val=0.3637, WAPE_test=0.4102).\n- Comparacion formal modelo vs plan humano: WAPE_modelo=0.3940 vs WAPE_planificada=0.0980.\n- Decision operativa: estrategia hibrida (plan humano base + alertas del modelo).",
        191: "CONCLUSIONES",
        192: "El proyecto alcanzo un nivel operativo y metodologico defendible para forecasting de produccion: pipeline E2E implementado, calidad de datos controlada, evaluacion temporal robusta y despliegue funcional en dashboard. Se cumplio el objetivo de transformar un enfoque inicial amplio en una solucion concentrada en una pregunta de negocio prioritaria.",
        194: "Trabajo futuro",
        195: "Como mejoras siguientes se proponen: (i) comparar contra plan ex-ante para evitar sesgo por ajustes ex-post, (ii) incorporar variables exogenas de negocio (promociones, quiebres, calendario comercial), (iii) monitorear drift de error por segmento, y (iv) evolucionar de alertas a politicas de recomendacion automatizadas por categoria/tipo.",
        197: "Limitaciones y consideraciones",
        198: "El corte actual evidencia que la planificacion humana esta altamente correlacionada con el real, por lo que el modelo no debe reemplazarla de forma directa en esta etapa. El valor inmediato del sistema es mejorar gobernanza, trazabilidad y capacidad de deteccion temprana de desalineaciones.",
        199: "REFERENCIAS BIBLIOGRAFICAS",
        201: "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",
        202: "Chapman, P., Clinton, J., Kerber, R., Khabaza, T., Reinartz, T., Shearer, C., & Wirth, R. (2000). CRISP-DM 1.0: Step-by-step data mining guide. SPSS Inc.",
        203: "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice (3rd ed.). OTexts.",
        204: "Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques (3rd ed.). Morgan Kaufmann.",
        205: "Jain, A. K. (2010). Data Clustering: 50 years beyond K-means. Pattern Recognition Letters, 31(8), 651-666.",
        206: "Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit (3rd ed.). Wiley.",
        207: "Mage AI. (2026). Documentation. https://docs.mage.ai/",
        208: "PostgreSQL Global Development Group. (2026). PostgreSQL Documentation. https://www.postgresql.org/docs/",
        209: "Streamlit Inc. (2026). Streamlit Documentation. https://docs.streamlit.io/",
        210: "Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). Data Mining: Practical Machine Learning Tools and Techniques (4th ed.). Morgan Kaufmann.",
        211: "Chandola, V., Banerjee, A., & Kumar, V. (2009). Anomaly detection: A survey. ACM Computing Surveys, 41(3), 1-58.",
        212: "",
        214: "ANEXO A: EVIDENCIAS DE PIPELINES Y CONSULTAS SQL",
        216: "Figura A1. Pipeline ETL Bronze en Mage. (captura: C06_mage_etl_bronze.png)\nFigura A2. Pipeline ETL Silver en Mage. (captura: C07_mage_etl_silver.png)\nFigura A3. Pipeline ETL Gold en Mage. (captura: C08_mage_etl_gold.png)\nFigura A4. Consulta muestra de Silver para modelado. (captura: C09_sql_silver_muestra.png)\nFigura A5. Conteos de Silver y distribucion PT/PP/OTRO. (captura: C10_sql_silver_conteos.png)\nFigura A6. Conteos de Gold v2 (activos/inactivos). (captura: C11_sql_gold_v2_conteos.png)\nFigura A7. Validacion de qty_planificada en Gold v2. (captura: C12_sql_gold_qty_planificada.png)\nFigura A8. Muestra de resultados operativos en Gold v2. (captura: C13_sql_gold_muestra.png).",
        218: "ANEXO B: EVIDENCIAS DE DASHBOARD OPERATIVO",
        220: "Figura B1. Resumen ejecutivo en Streamlit. (captura: C14_dashboard_resumen.png)\nFigura B2. KPIs de pronostico de produccion. (captura: C15_dashboard_predicciones_kpis.png)\nFigura B3. Top productos para planificacion. (captura: C16_dashboard_top_productos.png)\nFigura B4. Tabla de detalle con plan final hibrido. (captura: C17_dashboard_detalle_plan.png)\nFigura B5. Productos excluidos por vigencia. (captura: C18_dashboard_inactivos.png)\nFigura B6. Modelo publicado en ejecucion actual. (captura: C19_dashboard_modelo_publicado.png).",
        226: "ANEXO C: ARTEFACTOS Y REPORTES DE MODELADO",
        228: "Figura C1. Estructura del proyecto forecasting_tesis_v2. (captura: C01_estructura_proyecto.png)\nFigura C2. Artefactos generados por pipeline de modelado. (captura: C02_artefactos_generados.png)\nFigura C3. Reporte de wrangling y calidad de datos. (captura: C03_wrangling_report.png)\nFigura C4. Benchmark de modelos y baselines temporales. (captura: C04_benchmark_modelos.png)\nFigura C5. Reporte anti-leakage. (captura: C05_leakage_report.png)\nFigura C6. Documento tecnico E2E del proyecto. (captura: C20_md_documentacion_e2e.png).",
    }

    for idx, txt in repl.items():
        if idx < len(d.paragraphs):
            d.paragraphs[idx].text = txt

    for idx in [215, 217, 219, 221, 222, 223, 224, 225, 227, 229, 230]:
        if idx < len(d.paragraphs):
            d.paragraphs[idx].text = ""

    d.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
