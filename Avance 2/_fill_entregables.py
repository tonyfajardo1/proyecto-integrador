from docx import Document


BASE = r"F:\proyecto-integrador\Avance 2"


def fill_entregable_template():
    src = BASE + r"\plantilla_entregables (1).docx"
    doc = Document(src)

    replacements = {
        "Entregable X del Proyecto Integrador": "Entregable 2 del Proyecto Integrador",
        "Autor: Juan P\u00e9rez": "Autor: Anthony Fajardo",
        "202X": "2026",
        "T\u00edtulo del Proyecto:": "T\u00edtulo del Proyecto: Sistema de forecasting de produccion para soporte a planificacion operativa en CONDIMENSA",
        "Colocar el t\u00edtulo del proyecto. Por ejemplo: Aplicaci\u00f3n de algortimos de aprendizaje autom\u00e1tico para mejorar la provisi\u00f3n de calidad de servicio en redes definidas por software.": "Se implemento un flujo end-to-end de pronostico mensual de produccion: ingesta en Mage, consolidacion Silver, wrangling temporal con control de calidad, benchmark de modelos con validacion temporal, publicacion a Gold v2 e integracion con dashboard de planificacion.",
        "Resumen de actividades realizadas:": "Resumen de actividades realizadas (ultimo periodo):",
        "Enumerar las actividades realizadas durante el \u00faltimo per\u00edodo de trabajo y colocar una breve descripci\u00f3n o explicaci\u00f3n de la actividad. Por ejemplo:": "1) Se restructuro el alcance del proyecto hacia forecasting de produccion.\n2) Se creo el proyecto limpio 03_modelado/forecasting_tesis_v2 con notebooks y scripts reproducibles.\n3) Se realizo EDA y wrangling con controles de nulos, duplicados, estacionalidad e imputacion temporal/estructural.\n4) Se incorporaron controles anti-leakage (split temporal, encoding solo con train, reglas estacionales train+val).\n5) Se ejecuto benchmark (Baseline_Lag1, Lag12, RF, ExtraTrees, Prophet, ensemble y variantes segmentadas).\n6) Se publico resultado en gold.pronostico_produccion_resultado_v2 y se integro fallback en dashboard.\n7) Se restauro seccion de productos inactivos y estrategia hibrida plan humano + alertas del modelo.\n8) Se documentaron resultados, evidencias SQL/Streamlit y plan de mejora metodologica (ex-ante, segmentacion, metrica por clase minoritaria).",
        "Revisi\u00f3n del estado del arte: se realiz\u00f3 una b\u00fasqueda y selecci\u00f3n de art\u00edculos cient\u00edficos relevantes para identificar las propuestas actuales. De igual manera, se identificaron problemas o temas que no han sido abordados hasta el momento.": "Se actualizo el estado del arte enfocando trabajos relacionados en forecasting operacional, medallion architecture (OLTP/OLAP), y evaluacion temporal robusta para evitar data leakage.",
        "An\u00e1lisis de datasets: se realiz\u00f3 una b\u00fasqueda de datasets de red para utilizarlos con modelos de machine learning. Adicionalmente, se realiz\u00f3 un an\u00e1lisis de la estructura del dataset con el objetivo de determinar la calidad.": "Se audito el dataset de produccion (12,381 filas iniciales), con 0 nulos en campos clave, 0 duplicados producto-periodo y control de outliers sospechosos con trazabilidad.",
        "Secciones o cap\u00edtulos del documento final desarrollados": "Secciones o capitulos del documento final desarrollados en este periodo",
        "Enumerar las secciones o cap\u00edtulos del documento final que se han desarrollado durante el \u00faltimo per\u00edodo de trabajo.": "1) Introduccion y planteamiento del problema.\n2) Estado del arte (trabajos relacionados).\n3) Arquitectura y metodologia (integracion Mage-Silver-Gold-Streamlit, no en anexos).\n4) Desarrollo tecnico: ETL, wrangling, modelado y anti-leakage.\n5) Resultados experimentales y comparacion formal modelo vs plan humano.\n6) Conclusiones y lineas futuras.",
        "Introducci\u00f3n.": "Introduccion (redactada y actualizada).",
        "Estado del arte.": "Estado del arte (reestructurado con enfoque en trabajos relacionados).",
        "Propuesta.": "Propuesta tecnica (arquitectura medallion y pipeline E2E) desarrollada.",
        "Revisi\u00f3n y firma del tutor del proyecto": "Revision y firma del tutor del proyecto (pendiente de firma manuscrita)",
        "Yo, Ricardo Flores Moyano, profesor de la carrera de Ingenier\u00eda en Ciencias de la Computaci\u00f3n, hago constar que he revisado y, por lo tanto, apruebo las actividades realizadas durante este per\u00edodo de trabajo. Por otra parte, considero que el avance del proyecto integrador es adecuado y se corresponde con el cronograma definido en el documento de planificaci\u00f3n.": "Yo, Ricardo Flores Moyano, profesor de la carrera de Ingenieria en Ciencias de la Computacion, revisare este entregable para su validacion formal. Nota: para que el entregable sea valido en D2L, este documento debe incluir firma del tutor antes de su envio.",
        "Quito, X de XXXX de 202X": "Quito, 29 de marzo de 2026",
    }

    for paragraph in doc.paragraphs:
        if paragraph.text in replacements:
            paragraph.text = replacements[paragraph.text]

    out = BASE + r"\entregable_2_completado_avance2.docx"
    doc.save(out)
    return out


def fill_final_doc_template():
    src = BASE + r"\guia-biblioteca-pregrado (1).docx"
    doc = Document(src)

    replacements = {
        "Colegio de xxxx": "Colegio de Ciencias e Ingenierias",
        "T\u00edtulo del trabajo de la materia final de carrera": "Forecasting de produccion para soporte a planificacion operativa en CONDIMENSA",
        "T\u00edtulo del Trabajo de la materia final de carrera": "Forecasting de produccion para soporte a planificacion operativa en CONDIMENSA",
        "Nombre del estudiante (Completo)": "Anthony Fajardo",
        "Carrera": "Ingenieria en Ciencias de la Computacion",
        "XXXX": "Ingeniero en Ciencias de la Computacion",
        "Quito, d\u00eda de mes de a\u00f1o": "Quito, 29 de marzo de 2026",
        "Nombres y apellidos:                 xxxxxxx   xxxxx  xxxxx xxxxxxx": "Nombres y apellidos:                 Anthony Fajardo",
        "C\u00f3digo:                                         xxxxxxx": "Codigo:                                         [COMPLETAR CODIGO USFQ]",
        "C\u00e9dula de identidad:                  xxxxxxxxx": "Cedula de identidad:                  [COMPLETAR CEDULA]",
        "Lugar y fecha:\t\t\t   Ciudad, d\u00eda de mes de a\u00f1o": "Lugar y fecha:\t\t\t   Quito, 29 de marzo de 2026",
        "En texto normal, debes presentar una descripci\u00f3n completa pero concisa de tu trabajo, que motive a potenciales lectores a revisarlo por completo. El resumen debe indicar claramente cu\u00e1l es el asunto tratado en el trabajo, haciendo referencia a las motivaciones y enfoques utilizados para su desarrollo, los resultados m\u00e1s destacables, y las principales conclusiones que indiquen las implicaciones actuales y perspectivas futuras del asunto.": "Este trabajo presenta la implementacion de un sistema de forecasting mensual de produccion para CONDIMENSA, con el objetivo de apoyar la planificacion operativa mediante un flujo reproducible de datos y analitica. Se desarrollo una arquitectura Medallion (Bronze, Silver, Gold) usando Mage para ETL, una etapa de wrangling temporal con controles de calidad (nulos, duplicados, estacionalidad, imputacion y control de outliers sospechosos), y un benchmark de modelos con validacion temporal y controles anti-leakage. El mejor modelo por validacion fue RandomForest (WAPE_val ~0.364; WAPE_test ~0.410), mejorando frente a baselines temporales. Sin embargo, la comparacion formal contra la planificacion humana (qty_planificada) muestra que el plan actual mantiene menor error en este corte, por lo que se propone una estrategia hibrida: mantener el plan humano como base y usar el modelo como sistema de alertas y soporte a decision. Los resultados se publican en gold.pronostico_produccion_resultado_v2 y se visualizan en dashboard Streamlit con trazabilidad operativa.",
        "Palabras clave: Deben incluir entre 5 y 10 palabras claves que describan tu art\u00edculo.": "Palabras clave: forecasting, series de tiempo, arquitectura medallion, ETL, leakage, planificacion de produccion, dashboard, machine learning.",
        "En texto normal, debe ser una traducci\u00f3n precisa del resumen.": "This work presents the implementation of a monthly production forecasting system for CONDIMENSA to support operational planning through a reproducible data and analytics pipeline. A Medallion architecture (Bronze, Silver, Gold) was implemented using Mage for ETL, followed by time-series wrangling with data quality controls (nulls, duplicates, seasonality handling, imputation, and suspicious outlier control), and model benchmarking with temporal validation and anti-leakage checks. RandomForest achieved the best validation performance (WAPE_val ~0.364; WAPE_test ~0.410), outperforming temporal baselines. However, a formal comparison against human planning (qty_planificada) shows that current planning still yields lower error in this data cut. Therefore, a hybrid strategy is recommended: keep the human plan as the operational baseline and use the model as an alerting and decision-support layer. Outputs are published to gold.pronostico_produccion_resultado_v2 and consumed in a Streamlit dashboard.",
        "Key words: Presentar una traducci\u00f3n precisa de las palabras clave.": "Key words: forecasting, time series, Medallion architecture, ETL, leakage, production planning, dashboard, machine learning.",
        "Introducci\u00f3n\t10": "Introduccion\t10",
        "Desarrollo del Tema\t11": "Desarrollo del tema\t11",
        "Conclusiones\t12": "Conclusiones\t12",
        "Referencias bibliogr\u00e1ficas (ejemplo estilo APA)\t13": "Referencias bibliograficas\t13",
        "Anexo A: T\u00edtulo\t14": "Anexo A: Evidencias de pipeline y SQL\t14",
        "Anexo B: T\u00edtulo\t14": "Anexo B: Evidencias de dashboard\t15",
        "Anexo C: T\u00edtulo\t14": "Anexo C: Artefactos de modelado y reportes\t16",
        "Incluir una introducci\u00f3n en la que se explique de qu\u00e9 se trata el trabajo final, es decir una descripci\u00f3n general del tema, explicar de d\u00f3nde proviene el tema (diagn\u00f3stico), cu\u00e1l es su relevancia dentro del \u00e1rea del conocimiento y para el contexto ecuatoriano (antecedentes), si es necesario  en esta secci\u00f3n se debe incluir una definici\u00f3n de t\u00e9rminos, para terminar de describir los elementos que se encontrar\u00e1n en las siguientes secciones de tu trabajo final (al final de la Introducci\u00f3n debes tener unas oraciones de transici\u00f3n a la segunda parte).": "La planificacion de produccion en empresas de consumo masivo requiere decisiones mensuales oportunas, trazables y sustentadas en datos. En el contexto del proyecto integrador, se identifico la necesidad de unificar la gestion de datos operativos y analiticos para construir recomendaciones de produccion confiables. Por ello se adopto una arquitectura Medallion: OLTP (fuentes transaccionales) alimenta Bronze/Silver, mientras que OLAP consolida indicadores y resultados de prediccion en Gold para consumo en dashboard. Este trabajo se enfoca en forecasting mensual, con evaluacion temporal rigurosa para evitar data leakage y sobreajuste. En las secciones siguientes se describen trabajos relacionados, metodologia, resultados y limitaciones.",
        "En esta secci\u00f3n se desarrolla el tema elegido para el trabajo final. Se debe guardar el rigor acad\u00e9mico haciendo referencia a la bibliograf\u00eda utilizada. El estudiante deber\u00e1 desarrollar el tema en funci\u00f3n del entregable final que corresponde:": "El desarrollo integra cuatro componentes: (1) Ingenieria de datos con Mage y arquitectura Medallion para trazabilidad Bronze-Silver-Gold; (2) Preparacion del dataset de modelado mensual con control de calidad (nulos, duplicados, consistencia, estacionalidad, imputacion y control inteligente de outliers); (3) Modelado de series temporales con split temporal estricto, benchmark de baselines y modelos de machine learning, y controles anti-leakage (encoding de producto solo en train, reglas estacionales con train+val, cap por historico); y (4) Publicacion y visualizacion operativa en dashboard Streamlit.\n\nAdicionalmente, en respuesta a la retroalimentacion del primer avance, se fortalecio la organizacion del documento: la arquitectura se integra en metodologia (no en anexos), se incorporan trabajos relacionados en estado del arte, se justifica OLTP/OLAP para Medallion y se corrige el enfoque de evaluacion extrema de clasificacion (AUC-ROC=0.9993) migrando a metricas de forecasting y comparacion formal modelo vs plan humano.",
        "Un trabajo final o reporte escrito; o,": "El entregable corresponde a un trabajo final escrito con evidencia tecnica reproducible.",
        "La reflexi\u00f3n te\u00f3rica escrita de una presentaci\u00f3n o producci\u00f3n art\u00edstica abierta al p\u00fablico (carreras relacionadas con artes).": "No aplica para este proyecto (linea de ingenieria y analitica de datos).",
        "Presenta los aportes de este trabajo con base en lo investigado, es importante que como autor puedas analizar el tema y su relevancia para la profesi\u00f3n dentro del contexto nacional e internacional (presenta similitudes, diferencias entre los diferentes enfoques del tema investigado). En el caso de presentaciones art\u00edsticas o creativas se debe describir de qu\u00e9 se tratan y justificar sus elementos, obligatoriamente incluir anexos con fotos, evidencias (partituras, enlaces a videos, etc.) del producto elaborado. Realiza un an\u00e1lisis de lo que has aprendido en este trabajo, incluye sugerencias de estudios posibles que se realicen en el futuro para comprender de mejor manera el tema, menciona alguna dificultad que hayas tenido para realizar este trabajo y sus razones.": "El principal aporte es la implementacion de un pipeline de forecasting operacional defendible y reproducible para planificacion de produccion. A nivel metodologico, se evidencio que controles de leakage y validacion temporal son indispensables para evitar conclusiones optimistas. A nivel de negocio, el modelo mejora frente a baselines temporales, pero la planificacion humana actual mantiene mejor desempeno en este corte, por lo que se recomienda una estrategia hibrida: plan humano como base y modelo como sistema de alertas y apoyo para revision de SKU criticos.\n\nComo trabajo futuro se propone: i) separar plan ex-ante de ajustes ex-post para una comparacion justa, ii) incorporar variables exogenas (promociones, quiebres, calendario comercial), iii) evaluar precision, recall, F1 y AUPRC en tareas de clasificacion complementarias cuando corresponda, y iv) mantener repositorio versionado con evidencia de ejecucion y resultados.",
        "Referencias bibliogr\u00e1ficas": "Referencias bibliograficas",
        "Anexo A: T\u00edtulo": "Anexo A: Evidencias de pipelines Mage y consultas SQL",
        "Anexo B: T\u00edtulo": "Anexo B: Evidencias de dashboard operativo",
        "Anexo C: T\u00edtulo": "Anexo C: Artefactos de modelado y reportes",
    }

    for paragraph in doc.paragraphs:
        if paragraph.text in replacements:
            paragraph.text = replacements[paragraph.text]

    out = BASE + r"\documento_final_avance2_completado.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    out1 = fill_entregable_template()
    out2 = fill_final_doc_template()
    print(out1)
    print(out2)
