from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Abrir documento existente
doc = Document('CONDIMENSA_Proyecto_Data_Mining.docx')

doc.add_page_break()

# ============================================================================
# ARQUITECTURA MEDALLION
# ============================================================================
doc.add_heading('12. Arquitectura Medallion (Bronze/Silver/Gold)', 1)

doc.add_paragraph('''
La arquitectura Medallion organiza los datos en tres capas segun su nivel de procesamiento:

BRONZE (Datos Crudos):
- Datos tal como llegan de las fuentes
- Sin transformaciones
- Incluye metadata de carga (_loaded_at, _source)
- Sirve como respaldo y para reprocesamiento

SILVER (Datos Limpios):
- Datos validados y tipados correctamente
- Duplicados eliminados
- Valores nulos tratados
- Listos para analisis exploratorio

GOLD (Datos Analiticos):
- Modelo dimensional (Star Schema)
- Optimizado para consultas analiticas
- Metricas precalculadas
- Listo para dashboards y reportes
''')

doc.add_heading('12.1 Tablas por Capa', 2)

table = doc.add_table(rows=11, cols=3)
table.style = 'Table Grid'
data = [
    ['Capa', 'Tabla', 'Registros'],
    ['BRONZE', 'kronos_ventas_raw', '(staging)'],
    ['BRONZE', 'quickbooks_sales_raw', '(staging)'],
    ['BRONZE', 'quickbooks_items_raw', '(staging)'],
    ['SILVER', 'ventas', '1,000'],
    ['SILVER', 'agencias', '10'],
    ['SILVER', 'productos', '503'],
    ['GOLD', 'dim_tiempo', '1,096'],
    ['GOLD', 'dim_producto', '883'],
    ['GOLD', 'dim_agencia', '10'],
    ['GOLD', 'fact_ventas', '1,973'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

doc.add_page_break()

# ============================================================================
# STAR SCHEMA
# ============================================================================
doc.add_heading('13. Modelo Dimensional (Star Schema)', 1)

doc.add_paragraph('''
El modelo estrella (Star Schema) organiza los datos en:
- Tablas de Dimension: Describen las entidades (quien, que, cuando)
- Tabla de Hechos: Contiene las metricas/medidas del negocio

Ventajas:
- Consultas simples y rapidas (pocos JOINs)
- Facil de entender para usuarios de negocio
- Optimizado para herramientas de BI
- Permite drill-down y roll-up
''')

doc.add_heading('13.1 Dimension: dim_tiempo', 2)
code = '''CREATE TABLE gold.dim_tiempo (
    tiempo_id SERIAL PRIMARY KEY,
    fecha DATE NOT NULL UNIQUE,
    dia INTEGER,
    mes INTEGER,
    anio INTEGER,
    trimestre INTEGER,
    nombre_mes VARCHAR(20),
    nombre_dia VARCHAR(20),
    es_fin_semana BOOLEAN,
    semana_anio INTEGER
);'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('13.2 Dimension: dim_producto', 2)
code = '''CREATE TABLE gold.dim_producto (
    producto_id SERIAL PRIMARY KEY,
    codigo VARCHAR(100),
    nombre VARCHAR(200) NOT NULL,
    categoria VARCHAR(100),
    cluster_id INTEGER,
    cluster_nombre VARCHAR(100),
    nivel_riesgo VARCHAR(20),
    prob_devolucion NUMERIC(5,4)
);'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('13.3 Dimension: dim_agencia', 2)
code = '''CREATE TABLE gold.dim_agencia (
    agencia_id SERIAL PRIMARY KEY,
    codigo VARCHAR(100),
    nombre VARCHAR(200) NOT NULL,
    region VARCHAR(100),
    es_anomalia BOOLEAN DEFAULT FALSE,
    anomalia_score NUMERIC(10,4)
);'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('13.4 Tabla de Hechos: fact_ventas', 2)
code = '''CREATE TABLE gold.fact_ventas (
    venta_id SERIAL PRIMARY KEY,
    -- Foreign Keys (enlaces a dimensiones)
    tiempo_id INTEGER REFERENCES gold.dim_tiempo(tiempo_id),
    producto_id INTEGER REFERENCES gold.dim_producto(producto_id),
    agencia_id INTEGER REFERENCES gold.dim_agencia(agencia_id),
    -- Metricas/Medidas
    cantidad_venta NUMERIC(12,2),
    total_venta NUMERIC(14,2),
    cantidad_devolucion NUMERIC(12,2),
    total_devolucion NUMERIC(14,2),
    venta_neta NUMERIC(14,2),
    tasa_devolucion NUMERIC(5,2)
);

-- Indices para performance
CREATE INDEX idx_fact_ventas_tiempo ON gold.fact_ventas(tiempo_id);
CREATE INDEX idx_fact_ventas_producto ON gold.fact_ventas(producto_id);
CREATE INDEX idx_fact_ventas_agencia ON gold.fact_ventas(agencia_id);'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================================
# DIAGRAMA STAR SCHEMA
# ============================================================================
doc.add_heading('13.5 Diagrama del Star Schema', 2)

diagrama = '''
                              +------------------+
                              |   dim_tiempo     |
                              +------------------+
                              | tiempo_id (PK)   |
                              | fecha            |
                              | dia, mes, anio   |
                              | trimestre        |
                              | nombre_mes       |
                              +--------+---------+
                                       |
                                       |
+------------------+          +--------+---------+          +------------------+
|  dim_producto    |          |   fact_ventas    |          |   dim_agencia    |
+------------------+          +------------------+          +------------------+
| producto_id (PK) |<---------| venta_id (PK)    |--------->| agencia_id (PK)  |
| nombre           |          | tiempo_id (FK)   |          | nombre           |
| categoria        |          | producto_id (FK) |          | region           |
| cluster_nombre   |          | agencia_id (FK)  |          | es_anomalia      |
| nivel_riesgo     |          | cantidad_venta   |          | anomalia_score   |
| prob_devolucion  |          | total_venta      |          +------------------+
+------------------+          | total_devolucion |
                              | venta_neta       |
                              | tasa_devolucion  |
                              +------------------+
'''
p = doc.add_paragraph(diagrama)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_heading('13.6 Consultas de Ejemplo', 2)

doc.add_paragraph('Ventas por mes y agencia:')
code = '''SELECT
    t.nombre_mes,
    t.anio,
    a.nombre as agencia,
    SUM(f.total_venta) as ventas,
    SUM(f.total_devolucion) as devoluciones,
    AVG(f.tasa_devolucion) as tasa_promedio
FROM gold.fact_ventas f
JOIN gold.dim_tiempo t ON f.tiempo_id = t.tiempo_id
JOIN gold.dim_agencia a ON f.agencia_id = a.agencia_id
GROUP BY t.nombre_mes, t.anio, a.nombre
ORDER BY t.anio, t.mes;'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph('Productos de alto riesgo por cluster:')
code = '''SELECT
    p.cluster_nombre,
    p.nivel_riesgo,
    COUNT(*) as num_productos,
    SUM(f.total_venta) as ventas_totales,
    AVG(f.tasa_devolucion) as tasa_promedio
FROM gold.fact_ventas f
JOIN gold.dim_producto p ON f.producto_id = p.producto_id
WHERE p.nivel_riesgo = 'ALTO'
GROUP BY p.cluster_nombre, p.nivel_riesgo;'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# ============================================================================
# FLUJO COMPLETO
# ============================================================================
doc.add_heading('14. Flujo Completo de Datos', 1)

flujo = '''
+-------------+     +-------------+     +-------------+     +-------------+
|  FUENTES    |     |   BRONZE    |     |   SILVER    |     |    GOLD     |
|  EXTERNAS   |     |   (Raw)     |     |  (Clean)    |     | (Analytics) |
+-------------+     +-------------+     +-------------+     +-------------+
      |                   |                   |                   |
      |   ETL Pipeline    |   Transformacion  |   Modelo Dim.     |
      |   (Mage AI)       |   (Validacion)    |   (Star Schema)   |
      v                   v                   v                   v
+-------------+     +-------------+     +-------------+     +-------------+
| QuickBooks  |---->| qb_sales    |---->| ventas      |---->| fact_ventas |
| - sales     |     | qb_items    |     | productos   |     | dim_tiempo  |
| - items     |     |             |     | agencias    |     | dim_producto|
+-------------+     +-------------+     +-------------+     | dim_agencia |
                                                            +-------------+
+-------------+     +-------------+           |                   |
| Kronos      |---->| kr_ventas   |-----------+                   |
| - ventas    |     |             |                               |
+-------------+     +-------------+                               |
                                                                  v
                                                          +-------------+
                                                          | Dashboard   |
                                                          | Streamlit   |
                                                          | (Visuales)  |
                                                          +-------------+
'''
p = doc.add_paragraph(flujo)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_heading('14.1 Pipelines por Capa', 2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
data = [
    ['Capa', 'Pipeline', 'Entrada', 'Salida'],
    ['BRONZE', 'ETL_Kronos/QuickBooks', 'Supabase (raw)', 'bronze.*_raw'],
    ['SILVER', 'Transform_Clean', 'bronze.*', 'silver.ventas/productos/agencias'],
    ['GOLD', 'DM_Analytics', 'silver.*', 'gold.dim_*/fact_ventas'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

# Guardar
doc.save('CONDIMENSA_Proyecto_Data_Mining.docx')
print("Documento actualizado con arquitectura Medallion y Star Schema")
