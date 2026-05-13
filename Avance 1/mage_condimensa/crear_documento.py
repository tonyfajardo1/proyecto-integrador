from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# PORTADA
doc.add_paragraph()
title = doc.add_heading('PROYECTO CONDIMENSA', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Sistema de Data Mining con Mage AI')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Proyecto Integrador - Febrero 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 1. ARQUITECTURA
doc.add_heading('1. Arquitectura del Sistema', 1)
doc.add_paragraph('''
El sistema utiliza 4 contenedores Docker:

- Mage AI (Puerto 6789): Orquestador de pipelines ETL y ML
- PostgreSQL DWH (Puerto 5433): Data Warehouse local
- pgAdmin (Puerto 5050): Administrador de base de datos
- Streamlit Dashboard (Puerto 8501): Visualizacion interactiva

Fuentes de datos externas:
- QuickBooks (Supabase US-East): Ventas, produccion, items
- Kronos (Supabase US-West): Ventas consolidadas
''')

# 2. TECNOLOGIAS
doc.add_heading('2. Tecnologias Utilizadas', 1)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
data = [
    ['Componente', 'Tecnologia', 'Proposito'],
    ['Orquestador', 'Mage AI', 'Pipelines ETL y ML'],
    ['Contenedores', 'Docker + Docker Compose', 'Despliegue'],
    ['Data Warehouse', 'PostgreSQL 15', 'Almacenamiento'],
    ['Dashboard', 'Streamlit + Plotly', 'Visualizacion'],
    ['Admin BD', 'pgAdmin 4', 'Administracion'],
    ['ML', 'Scikit-learn, mlxtend', 'Algoritmos ML'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

doc.add_page_break()

# 3. DOCKER COMPOSE
doc.add_heading('3. Docker Compose', 1)
docker_code = '''version: '3.8'

services:
  mage:
    image: mageai/mageai:latest
    container_name: mage_condimensa
    command: mage start condimensa_project
    environment:
      - QUICKBOOKS_HOST=your-quickbooks-host.supabase.com
      - QUICKBOOKS_PORT=6543
      - QUICKBOOKS_DB=postgres
      - QUICKBOOKS_USER=postgres.your-quickbooks-project-ref
      - QUICKBOOKS_PASSWORD=change_me
      - QUICKBOOKS_SCHEMA=quickbooks
      - KRONOS_HOST=your-kronos-host.supabase.com
      - KRONOS_PORT=6543
      - KRONOS_DB=postgres
      - KRONOS_USER=postgres.your-kronos-project-ref
      - KRONOS_PASSWORD=change_me
      - KRONOS_SCHEMA=kronos
    ports:
      - "6789:6789"
    volumes:
      - ./condimensa_project:/home/src/condimensa_project

  postgres_local:
    image: postgres:15
    container_name: condimensa_dwh
    environment:
      - POSTGRES_USER=condimensa
      - POSTGRES_PASSWORD=change_me
      - POSTGRES_DB=condimensa_analytics
    ports:
      - "5433:5432"

  pgadmin:
    image: dpage/pgadmin4:latest
    container_name: condimensa_pgadmin
    environment:
      - PGADMIN_DEFAULT_EMAIL=admin@condimensa.com
      - PGADMIN_DEFAULT_PASSWORD=change_me
    ports:
      - "5050:80"

  dashboard:
    image: python:3.11-slim
    container_name: condimensa_dashboard
    command: bash -c "pip install streamlit pandas plotly psycopg2-binary && streamlit run app.py"
    ports:
      - "8501:8501"
    volumes:
      - ./dashboard:/app
'''
p = doc.add_paragraph(docker_code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

doc.add_page_break()

# 4. PIPELINES
doc.add_heading('4. Pipelines Creados (7 Total)', 1)

doc.add_heading('4.1 Pipelines ETL', 2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
data = [
    ['Pipeline', 'Trigger', 'Descripcion'],
    ['01_ETL_Kronos_Ventas', 'Diario 6:00 AM', 'Extrae ventas de Kronos'],
    ['02_ETL_QuickBooks_Ventas', 'Diario 6:15 AM', 'Extrae ventas de QuickBooks'],
    ['03_ETL_QuickBooks_Produccion', 'Diario 6:30 AM', 'Extrae produccion'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

doc.add_paragraph()
doc.add_heading('4.2 Pipelines Data Mining', 2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
data = [
    ['Pipeline', 'Algoritmo', 'Trigger', 'Descripcion'],
    ['04_DM_Deteccion_Anomalias', 'Isolation Forest', 'Lunes 7:00', 'Detecta agencias atipicas'],
    ['05_DM_Reglas_Asociacion', 'Apriori', 'Lunes 7:30', 'Patrones de productos'],
    ['06_DM_Prediccion_Devoluciones', 'Random Forest', 'Lunes 8:00', 'Predice riesgo'],
    ['07_DM_Clustering_Productos', 'K-Means', 'Lunes 8:30', 'Segmenta productos'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

doc.add_page_break()

# 5. CODIGO ISOLATION FOREST
doc.add_heading('5. Codigo: Deteccion de Anomalias (Isolation Forest)', 1)
doc.add_paragraph('Isolation Forest aisla observaciones anomalas. Las anomalias son mas faciles de aislar porque tienen valores muy diferentes.')
code = '''from sklearn.ensemble import IsolationForest
from sklearn.preprocessing import StandardScaler

@transformer
def detect_anomalies(df, *args, **kwargs):
    # Seleccionar features
    features = ['total_ventas', 'total_devoluciones', 'ratio_devolucion']
    X = df[features].copy()

    # Escalar datos
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # Entrenar modelo
    model = IsolationForest(
        contamination=0.1,   # 10% anomalias esperadas
        random_state=42,
        n_estimators=100
    )

    # Predecir: -1 = anomalia, 1 = normal
    predictions = model.fit_predict(X_scaled)

    df['anomalia_score'] = model.decision_function(X_scaled)
    df['es_anomalia'] = predictions == -1

    return df
'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# 6. CODIGO APRIORI
doc.add_heading('6. Codigo: Reglas de Asociacion (Apriori)', 1)
doc.add_paragraph('''
Apriori encuentra patrones frecuentes en datos transaccionales.
- Support: Frecuencia de aparicion
- Confidence: P(B|A) - Probabilidad de B dado A
- Lift: Cuanto aumenta la probabilidad de B cuando A esta presente
''')
code = '''from mlxtend.frequent_patterns import apriori, association_rules

@transformer
def generate_association_rules(df_binary, *args, **kwargs):
    # Encontrar itemsets frecuentes (min 5% de transacciones)
    frequent_itemsets = apriori(
        df_binary,
        min_support=0.05,
        use_colnames=True
    )

    # Generar reglas con lift > 1
    rules = association_rules(
        frequent_itemsets,
        metric="lift",
        min_threshold=1.0
    )

    # Filtrar por confianza alta
    rules = rules[rules['confidence'] >= 0.5]
    rules = rules.sort_values('lift', ascending=False)

    return rules
'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# 7. CODIGO RANDOM FOREST
doc.add_heading('7. Codigo: Prediccion (Random Forest)', 1)
doc.add_paragraph('Random Forest combina multiples arboles de decision para predicciones robustas.')
code = '''from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score

@transformer
def train_prediction_model(df, *args, **kwargs):
    # Features y target
    X = df[['cant_venta', 'total_venta', 'ratio_cantidad']]
    y = (df['tasa_devolucion'] > 15).astype(int)  # 1 = devolucion alta

    # Split train/test
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2)

    # Entrenar modelo
    model = RandomForestClassifier(
        n_estimators=100,
        max_depth=10,
        random_state=42,
        class_weight='balanced'
    )
    model.fit(X_train, y_train)

    # Evaluar
    accuracy = accuracy_score(y_test, model.predict(X_test))
    print(f"Precision: {accuracy:.2%}")

    # Predecir probabilidades
    df['prob_devolucion_alta'] = model.predict_proba(X)[:, 1]

    # Clasificar riesgo
    df['nivel_riesgo'] = df['prob_devolucion_alta'].apply(
        lambda p: 'ALTO' if p >= 0.7 else ('MEDIO' if p >= 0.4 else 'BAJO')
    )

    return df
'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# 8. CODIGO K-MEANS
doc.add_heading('8. Codigo: Clustering (K-Means)', 1)
doc.add_paragraph('K-Means agrupa datos en K clusters basandose en distancia a centroides.')
code = '''from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler

@transformer
def apply_kmeans_clustering(df, *args, **kwargs):
    # Features para clustering
    X = df[['total_ventas', 'tasa_devolucion']].copy()

    # Escalar (K-Means es sensible a escala)
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)

    # Aplicar K-Means con 4 clusters
    kmeans = KMeans(n_clusters=4, random_state=42, n_init=10)
    df['cluster'] = kmeans.fit_predict(X_scaled)

    # Nombrar clusters segun centroides
    centroids = scaler.inverse_transform(kmeans.cluster_centers_)

    cluster_names = {}
    for i, (ventas, dev) in enumerate(centroids):
        vol = "Alto Vol" if ventas > df['total_ventas'].median() else "Bajo Vol"
        dv = "Alta Dev" if dev > df['tasa_devolucion'].median() else "Baja Dev"
        cluster_names[i] = f"{vol} - {dv}"

    df['cluster_nombre'] = df['cluster'].map(cluster_names)

    return df
'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_page_break()

# 9. RESULTADOS
doc.add_heading('9. Resultados Obtenidos', 1)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
data = [
    ['Metrica', 'Valor'],
    ['Agencias analizadas', '10'],
    ['Anomalias detectadas', '2 (Santo Domingo, Ibarra)'],
    ['Productos segmentados', '430'],
    ['Clusters identificados', '4'],
    ['Precision modelo predictivo', '80%'],
    ['Productos alto riesgo', '~150'],
    ['Tasa devolucion global', '~14%'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

doc.add_paragraph()
doc.add_heading('Interpretacion de Clusters:', 2)
doc.add_paragraph('''
Cluster 0: Bajo Volumen - Baja Devolucion
  - Productos estables con pocas ventas
  - Mantener inventario minimo

Cluster 1: Alto Volumen - Baja Devolucion
  - Productos estrella
  - Maximizar disponibilidad

Cluster 2: Bajo Volumen - Alta Devolucion
  - Productos problematicos
  - Investigar causas

Cluster 3: Alto Volumen - Alta Devolucion
  - Productos populares con problemas
  - Prioridad para mejora de calidad
''')

doc.add_page_break()

# 10. COMANDOS
doc.add_heading('10. Comandos Utiles', 1)
code = '''# Levantar servicios
docker-compose up -d

# Ver contenedores
docker ps

# Ver logs
docker logs mage_condimensa
docker logs condimensa_dashboard

# Reiniciar servicio
docker-compose restart dashboard

# Detener todo
docker-compose down

# Conectar a PostgreSQL
docker exec -it condimensa_dwh psql -U condimensa -d condimensa_analytics

# Ver tablas
\\dt

# Contar registros
SELECT COUNT(*) FROM dm_anomalias_agencias;
'''
p = doc.add_paragraph(code)
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

# 11. URLs
doc.add_heading('11. URLs de Acceso', 1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
data = [
    ['Servicio', 'URL', 'Credenciales'],
    ['Mage AI', 'http://localhost:6789', '-'],
    ['Dashboard', 'http://localhost:8501', '-'],
    ['pgAdmin', 'http://localhost:5050', 'admin@condimensa.com / change_me'],
    ['PostgreSQL', 'localhost:5433', 'condimensa / change_me'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = cell

# GUARDAR
doc.save('CONDIMENSA_Proyecto_Data_Mining.docx')
print("Documento creado: CONDIMENSA_Proyecto_Data_Mining.docx")
